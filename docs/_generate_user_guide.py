#!/usr/bin/env python3
"""Generate docs/user-guide.docx and docs/user-guide.html from captured screenshots."""

from html import escape
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "user-guide-assets"
OUTPUT = ROOT / "user-guide.docx"
OUTPUT_HTML = ROOT / "user-guide.html"

NAVY = RGBColor(0x0F, 0x17, 0x2A)
BLUE = RGBColor(0x1D, 0x4E, 0xD8)
MUTED = RGBColor(0x47, 0x55, 0x69)
AMBER = RGBColor(0xB4, 0x53, 0x09)


class HtmlSink:
    def __init__(self):
        self.body = []
        self.toc = []
        self._slugs = set()

    def slug(self, text):
        base = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-") or "section"
        slug = base
        n = 2
        while slug in self._slugs:
            slug = f"{base}-{n}"
            n += 1
        self._slugs.add(slug)
        return slug

    def add(self, html):
        self.body.append(html)

    def heading(self, text, level):
        slug = self.slug(text)
        if level <= 2:
            self.toc.append((level, slug, text))
        self.add(f'<h{level} id="{slug}">{escape(text)}</h{level}>')

    def render(self):
        toc_items = []
        for level, slug, text in self.toc:
            cls = "toc-h1" if level == 1 else "toc-h2"
            toc_items.append(
                f'<a class="{cls}" href="#{slug}">{escape(text)}</a>'
            )
        return HTML_PAGE.format(
            toc="\n".join(toc_items),
            body="\n".join(self.body),
        )


HTML = HtmlSink()

HTML_PAGE = """<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>eSariSari User Guide</title>
  <style>
    :root {{
      --navy: #0f172a;
      --blue: #1d4ed8;
      --muted: #475569;
      --line: #e2e8f0;
      --bg: #f8fafc;
    }}
    * {{ box-sizing: border-box; }}
    html {{ scroll-behavior: smooth; }}
    body {{
      margin: 0;
      font-family: Calibri, "Segoe UI", system-ui, sans-serif;
      color: var(--navy);
      background: var(--bg);
      line-height: 1.5;
    }}
    .wrap {{
      display: grid;
      grid-template-columns: 16rem minmax(0, 48rem);
      gap: 2rem;
      max-width: 72rem;
      margin: 0 auto;
      padding: 2rem 1.25rem 4rem;
    }}
    nav {{
      position: sticky;
      top: 1rem;
      align-self: start;
      max-height: calc(100vh - 2rem);
      overflow: auto;
      padding-right: 0.5rem;
    }}
    nav .label {{
      font-size: 0.7rem;
      font-weight: 700;
      letter-spacing: 0.08em;
      text-transform: uppercase;
      color: var(--muted);
      margin-bottom: 0.75rem;
    }}
    nav a {{
      display: block;
      color: var(--muted);
      text-decoration: none;
      font-size: 0.85rem;
      padding: 0.2rem 0;
    }}
    nav a:hover {{ color: var(--blue); }}
    nav .toc-h1 {{ font-weight: 700; color: var(--navy); margin-top: 0.6rem; }}
    nav .toc-h2 {{ padding-left: 0.75rem; font-size: 0.8rem; }}
    article {{
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 1rem;
      padding: 2rem 1.75rem 3rem;
    }}
    .kicker {{ color: var(--blue); font-weight: 700; margin: 0; }}
    .lede {{ color: var(--muted); margin: 0 0 0.5rem; }}
    h1.cover {{ font-size: 2.25rem; margin: 0.25rem 0 0.5rem; }}
    h1 {{ font-size: 1.55rem; margin: 2rem 0 0.6rem; }}
    h2 {{ font-size: 1.2rem; margin: 1.5rem 0 0.5rem; }}
    h3 {{ font-size: 1.05rem; margin: 1.25rem 0 0.4rem; }}
    p {{ margin: 0 0 0.85rem; }}
    p.muted {{ color: var(--muted); font-style: italic; }}
    ol, ul {{ margin: 0 0 1rem; padding-left: 1.25rem; }}
    li {{ margin: 0.25rem 0; }}
    ol li::marker {{ color: var(--blue); font-weight: 700; }}
    table {{
      width: 100%;
      border-collapse: collapse;
      font-size: 0.9rem;
      margin: 0 0 1.1rem;
    }}
    th, td {{
      border: 1px solid var(--line);
      padding: 0.45rem 0.55rem;
      text-align: left;
      vertical-align: top;
    }}
    th {{ background: var(--navy); color: #fff; font-size: 0.78rem; }}
    tr:nth-child(even) td {{ background: #f8fafc; }}
    .callout {{
      border: 1px solid #c7d2fe;
      border-radius: 0.75rem;
      padding: 0.85rem 1rem;
      margin: 0 0 1rem;
    }}
    .callout strong {{ color: var(--blue); display: block; margin-bottom: 0.25rem; }}
    figure {{ margin: 1rem 0 1.25rem; text-align: center; }}
    figure img {{
      max-width: 100%;
      height: auto;
      border: 1px solid var(--line);
      border-radius: 0.5rem;
    }}
    figcaption {{ color: var(--muted); font-size: 0.85rem; font-style: italic; margin-top: 0.4rem; }}
    .missing {{ color: #b45309; font-style: italic; }}
    .download {{ font-size: 0.9rem; margin: 0 0 1.25rem; }}
    .download a {{ color: var(--blue); }}
    @media (max-width: 860px) {{
      .wrap {{ grid-template-columns: 1fr; }}
      nav {{ position: static; max-height: none; }}
    }}
    @media print {{
      nav {{ display: none; }}
      .wrap {{ display: block; max-width: none; padding: 0; }}
      article {{ border: 0; }}
    }}
  </style>
</head>
<body>
  <div class="wrap">
    <nav>
      <div class="label">Contents</div>
      {toc}
    </nav>
    <article>
      {body}
    </article>
  </div>
</body>
</html>
"""


def set_run(run, *, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="CBD5E1"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_para(doc, text, *, size=11, bold=False, color=None, space_after=8, space_before=0, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color, italic=italic)
    if size >= 28:
        HTML.add(f'<h1 class="cover">{escape(text)}</h1>')
    elif size >= 14 and color == BLUE:
        weight = "700" if bold else "600"
        HTML.add(f'<p class="kicker" style="font-weight:{weight}">{escape(text)}</p>')
    elif size >= 11 and color == MUTED and not italic:
        HTML.add(f'<p class="lede">{escape(text)}</p>')
    else:
        cls = ' class="muted"' if italic or color == MUTED else ""
        HTML.add(f"<p{cls}>{escape(text)}</p>")
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(16 if level > 1 else 8)
    p.paragraph_format.space_after = Pt(8)
    HTML.heading(text, 3 if level >= 3 else level)
    return p


def shot(doc, filename, caption):
    path = ASSETS / filename
    if not path.exists():
        add_para(doc, f"[Screenshot missing: {filename}]", italic=True, color=AMBER)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.3))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    set_run(r, size=9, italic=True, color=MUTED)
    HTML.add(
        "<figure>"
        f'<img src="user-guide-assets/{escape(filename)}" alt="{escape(caption)}" />'
        f"<figcaption>{escape(caption)}</figcaption>"
        "</figure>"
    )


def steps(doc, items):
    lis = []
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        n = p.add_run(f"{i}.  ")
        set_run(n, size=11, bold=True, color=BLUE)
        t = p.add_run(item)
        set_run(t, size=11, color=NAVY)
        lis.append(f"<li>{escape(item)}</li>")
    HTML.add("<ol>" + "".join(lis) + "</ol>")


def bullets(doc, items):
    lis = []
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(3)
        n = p.add_run("•  ")
        set_run(n, size=11, color=BLUE)
        t = p.add_run(item)
        set_run(t, size=11, color=NAVY)
        lis.append(f"<li>{escape(item)}</li>")
    HTML.add("<ul>" + "".join(lis) + "</ul>")


def callout(doc, title, body, fill="EEF2FF"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_borders(cell, "C7D2FE")
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(title)
    set_run(r1, size=10, bold=True, color=BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(body)
    set_run(r2, size=10, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    HTML.add(
        f'<div class="callout" style="background:#{fill.lower()}">'
        f"<strong>{escape(title)}</strong>{escape(body)}</div>"
    )


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "0F172A")
        set_cell_borders(cell, "0F172A")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run(r, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for ri, row in enumerate(rows):
        for ci, value in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            shade_cell(cell, "F8FAFC" if ri % 2 == 0 else "FFFFFF")
            set_cell_borders(cell, "E2E8F0")
            p = cell.paragraphs[0]
            r = p.add_run(str(value))
            set_run(r, size=9, color=NAVY)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)
    head = "".join(f"<th>{escape(h)}</th>" for h in headers)
    body = "".join(
        "<tr>" + "".join(f"<td>{escape(str(value))}</td>" for value in row) + "</tr>"
        for row in rows
    )
    HTML.add(f"<table><thead><tr>{head}</tr></thead><tbody>{body}</tbody></table>")


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    core = doc.core_properties
    core.title = "eSariSari User Guide"
    core.subject = "Sub-Franchisee, Franchisee, and Retailer"
    core.author = "eSariSari"

    add_para(doc, "eSariSari", size=14, bold=True, color=BLUE, space_after=0)
    add_para(doc, "Franchise financial operations platform", size=11, color=MUTED, space_after=18)
    add_para(doc, "User Guide", size=28, bold=True, color=NAVY, space_after=4)
    add_para(
        doc,
        "Sub-Franchisee · Franchisee · Retailer",
        size=14,
        color=BLUE,
        space_after=12,
    )
    add_para(
        doc,
        "A click-by-click walkthrough of how to buy Internet Credits, release them down the chain (request or Direct Release), record a customer sale, and read both earnings streams on Revenue and Reports. Admin is not covered in this version. Detailed formulas are in section 1 and in the Excel workbook esari-computations.xlsx.",
        size=11,
        color=NAVY,
        space_after=10,
    )
    add_para(doc, "1 September 2026  ·  Demo environment", size=10, italic=True, color=MUTED)
    HTML.add(
        '<p class="download"><a href="user-guide.docx">Download User Guide Word copy</a> · <a href="esari-computations.xlsx" download="esari-computations.xlsx">Download computations (Excel)</a></p>'
    )

    callout(
        doc,
        "About this demo (version 1)",
        "This is the first walkthrough of the franchise operations demo. Sales and Internet Credits history start empty. Franchisees and retailers start at ₱0 Available Credits — restock with a request or Direct Release. There are two earnings streams: Internet Credits earnings and Sales Commission (from customer sales). The bell alerts the upline to pending downline credit requests and to low or zero balances. Screenshots show layout; numbers change as you walk the demo.",
        fill="EEF2FF",
    )

    heading(doc, "1. How the money works (read this first)", 1)
    add_para(
        doc,
        "Every role in this guide is part of one chain. Credits move down. Cash for loads moves up. Customer cash for internet sales is split as a % of the customer payment (Sales).",
    )

    add_table(
        doc,
        ["Who", "Buys credits from", "Sells / uses credits", "Typical deposit rate"],
        [
            ["Sub-Franchisee", "Admin / platform", "Franchisees", "Buys at 60%, sells at 70%"],
            ["Franchisee", "Sub-Franchisee", "Retailers", "Buys at 70%, sells at 80%"],
            ["Retailer", "Franchisee", "End customers (internet sales)", "Buys at 80%"],
        ],
    )

    heading(doc, "Two different earning streams", 2)
    add_para(
        doc,
        "Do not mix these up. The app prices them on different screens, and they show up as different numbers on Revenue and Reports.",
    )
    add_table(
        doc,
        ["Stream", "When it happens", "What you earn", "Where you set it"],
        [
            [
                "Internet Credits earnings",
                "A downline deposits cash and you release credits (request or Direct Release)",
                "Admin: cash collected. Sub/Fran: cash in minus what those credits cost you",
                "Deposit Rates",
            ],
            [
                "Sales Commission",
                "A retailer sells internet to a customer",
                "Your % of the customer payment (Sales)",
                "Commission Settings (Sub-Franchisee / Admin)",
            ],
        ],
    )

    heading(doc, "Credit loads — cash in, credits out", 2)
    add_para(
        doc,
        "A deposit rate is the cash paid divided by the credit face value. Formula: credits = deposit ÷ rate. Example at 70%: ₱7,000 cash buys 10,000 credits.",
    )
    add_para(
        doc,
        "Worked example on 1,000 credits (1 credit face = ₱1 for display):",
    )
    add_table(
        doc,
        ["Hop", "Cash paid", "Credits received", "Spread for the seller"],
        [
            ["Admin → Sub-Franchisee @ 60%", "₱600", "1,000", "Admin collects load cash"],
            ["Sub-Franchisee → Franchisee @ 70%", "₱700", "1,000", "Sub earns ₱100 Internet Credits earnings"],
            ["Franchisee → Retailer @ 80%", "₱800", "1,000", "Franchisee earns ₱100 Internet Credits earnings"],
        ],
    )
    callout(
        doc,
        "Cash vs credits",
        "Available Credits are inventory, not cash in the bank. When you release credits, your inventory goes down. The cash from the downline’s deposit stays with you as the collector for that hop. Always match proof of payment before you release. Direct Release uses the same formula and proof rules — it just skips the pending queue.",
    )

    heading(doc, "Internet sales — customer payment, then a split", 2)
    add_para(
        doc,
        "In this demo, the retailer records a sale from Transactions → Record demo sale. Credits consumed equal the customer payment (100% of sales). Commission Settings split that same customer payment (Sales).",
    )
    bullets(
        doc,
        [
            "Customer payment (Sales) — what the customer paid (example ₱1,000). Commission Settings % apply here.",
            "Credits consumed — inventory burned from the retailer’s Available Credits (example ₱1,000, equal to the payment).",
        ],
    )
    add_para(
        doc,
        "After a ₱1,000 demo sale on Retailer A (demo settings 10% / 20% / 30% / 40% of sales):",
    )
    add_table(
        doc,
        ["Party", "Share of ₱1,000 sales", "Amount"],
        [
            ["Retailer A", "10%", "₱100.00"],
            ["Franchisee A", "20%", "₱200.00"],
            ["Northern Mindanao Sub-Franchisee", "30%", "₱300.00"],
            ["eSariSari platform", "40%", "₱400.00"],
            ["Total", "100%", "₱1,000.00"],
        ],
    )
    callout(
        doc,
        "Credits consumed equal the payment",
        "If the customer paid ₱1,000, ₱1,000 of Available Credits leave the retailer wallet. The retailer’s earning is still 10% of the ₱1,000 (₱100) — commission is a split of sales, not leftover inventory. Read Your Commission on Revenue.",
        fill="ECFDF5",
    )

    heading(doc, "Detailed computations", 2)
    add_para(
        doc,
        "The app always rounds peso amounts to two decimal places (ROUND to 0.01). 1 credit face = ₱1 for display. Yellow cells in the Excel workbook esari-computations.xlsx are inputs; grey cells recompute. Direct Release uses the same load formulas as a released request.",
    )
    add_para(
        doc,
        "Keep the Excel file next to this guide (docs/esari-computations.xlsx), or download it from Sign In (download Excel) or from your name menu → Download computations (Excel). Sheets, left to right: Load calculator, Sale calculator, Worked examples, Demo defaults, then Read me and All formulas.",
    )

    heading(doc, "Internet Credits loads", 3)
    add_para(doc, "Deposit rate = cash paid ÷ credit face. Never add a percent to the cash.")
    add_table(
        doc,
        ["What", "Formula", "Worked example"],
        [
            ["Suggested credits", "credits = ROUND(deposit ÷ rate, 2)", "₱7,000 ÷ 70% = 10,000 credits"],
            ["Cash for 1,000 credits", "cash = ROUND(1,000 × rate, 2)", "1,000 × 80% = ₱800"],
            ["Admin IC earnings", "cash collected on released loads", "₱600 cash for 1,000 credits at 60%"],
            [
                "Sub / Fran IC earnings",
                "spread = cash in − ROUND(credits × your buy rate, 2)",
                "₱700 − (1,000 × 60%) = ₱100",
            ],
        ],
    )
    add_para(doc, "Default hops and what the seller earns on 1,000 credits:")
    add_table(
        doc,
        ["Hop", "Rate", "Cash paid", "Seller cost", "Seller IC earnings"],
        [
            ["Admin → Sub-Franchisee", "60%", "₱600", "— (Admin collects cash)", "₱600 cash collected"],
            ["Sub-Franchisee → Franchisee", "70%", "₱700", "₱600 (bought at 60%)", "₱100 spread"],
            ["Franchisee → Retailer", "80%", "₱800", "₱700 (bought at 70%)", "₱100 spread"],
            ["Retailer → customer", "n/a (sale, not a load)", "customer payment", "credits consumed", "No load spread — commission only"],
        ],
    )
    add_para(
        doc,
        "Guide example — Sub New Credits Request: ₱30,000 at 60% → 50,000 credits. Direct Release ₱7,000 at 70% → 10,000 credits; Sub cost ₱6,000; Sub IC earnings ₱1,000. Retailer request ₱1,600 at 80% → 2,000 credits.",
    )

    heading(doc, "Demo internet sale", 3)
    add_para(
        doc,
        "Record demo sale burns Available Credits at 100% of the customer payment. If credits required are more than Available Credits, the sale is blocked.",
    )
    add_table(
        doc,
        ["Line", "Formula", "On a ₱1,000 sale"],
        [
            ["Credits consumed", "ROUND(payment, 2)", "₱1,000.00"],
            ["Commission base (Sales)", "customer payment", "₱1,000.00"],
        ],
    )

    heading(doc, "Commission split", 3)
    add_para(
        doc,
        "Each party’s commission is ROUND(customer payment × their % ÷ 100, 2). Platform gets the remainder so the four shares always add up to the customer payment (leftover cents). Credits consumed stay inventory and are not subtracted from the split.",
    )
    add_para(
        doc,
        "On Commission Settings: Admin can edit the platform fee; Sub-Franchisee share is the remainder. A Sub-Franchisee can edit Retailer %, Franchisee %, and Your Share; the platform fee is Admin-set and read-only. The four percentages must total 100% of sales.",
    )
    add_table(
        doc,
        ["Retailer (demo seed)", "Retailer", "Franchisee", "Sub-Franchisee", "Platform"],
        [
            ["Retailer A / B / C", "10%", "20%", "30%", "40%"],
            ["New settings dialog default", "10%", "20%", "30%", "40%"],
        ],
    )
    add_table(
        doc,
        ["Party on a ₱1,000 Retailer A sale", "% of ₱1,000 sales", "Amount"],
        [
            ["Retailer A", "10%", "₱100.00"],
            ["Franchisee A", "20%", "₱200.00"],
            ["Northern Mindanao Sub-Franchisee", "30%", "₱300.00"],
            ["eSariSari platform (remainder)", "40%", "₱400.00"],
            ["Total", "100%", "₱1,000.00"],
        ],
    )
    add_para(
        doc,
        "Sub / Franchisee Revenue Total earnings = Internet Credits earnings (spreads in the period) + Sales Commission (credited shares in the period). Retailers have no load-spread stream — only their commission, sales volume, and credits consumed.",
    )

    heading(doc, "Available Credits and alerts", 3)
    add_table(
        doc,
        ["What", "Formula", "Example"],
        [
            [
                "Available Credits",
                "opening + credits received − credits released − credits consumed on sales",
                "Retailer A: 6,250 received − 970 consumed = ₱5,280",
            ],
            ["Sufficient", "available > ₱5,000", "₱5,280"],
            ["Low Balance", "₱0 < available ≤ ₱5,000", "Retailer B ₱2,500 — amber banner"],
            ["Zero Balance", "available ≤ ₱0", "Fresh franchisee / retailer — red banner"],
        ],
    )

    heading(doc, "2. Sign in", 1)
    add_para(doc, "Open the app and use the Sign In card. Demo accounts sit under the form.")
    shot(doc, "login.png", "Figure 1. Sign In screen with demo accounts.")
    steps(
        doc,
        [
            "Open the app: https://prlusterio.github.io/esari-finops-mvp/login",
            "Optional: at the bottom of Sign In, open the user guide, download User Guide Word copy, or download Excel (the computations workbook). After you sign in, the same Excel file is also under your name → Download computations (Excel).",
            "Under Demo Accounts, click Sub-Franchisee, Franchisee A, or Retailer A. That fills the email only — not the password.",
            "Type the shared demo password: password123. (Admin is marked Restricted and is not used in this guide.)",
            "Click Sign In.",
            "You land on that role’s home page: Wallets for Sub-Franchisee and Franchisee, Wallet for Retailer.",
        ],
    )
    add_table(
        doc,
        ["Role", "Demo email", "Password", "Home screen", "Opening Available Credits"],
        [
            [
                "Sub-Franchisee",
                "subfranchisee@esarisari.local",
                "password123",
                "Wallets",
                "₱135,000",
            ],
            [
                "Franchisee A",
                "franchisee-a@esarisari.local",
                "password123",
                "Wallets",
                "₱0 — request or receive a Direct Release",
            ],
            [
                "Franchisee B",
                "franchisee-b@esarisari.local",
                "password123",
                "Wallets",
                "₱0",
            ],
            [
                "Retailer A (under Fran A)",
                "retailer-a@esarisari.local",
                "password123",
                "Wallet",
                "₱0",
            ],
            [
                "Retailer B (under Fran A)",
                "retailer-b@esarisari.local",
                "password123",
                "Wallet",
                "₱0",
            ],
            [
                "Retailer C (under Fran B)",
                "retailer-c@esarisari.local",
                "password123",
                "Wallet",
                "₱0",
            ],
        ],
    )
    add_para(
        doc,
        "Use Franchisee A + Retailer A for the main walkthrough (same chain). The bell in the header shows pending Internet Credits requests from your downlines and low/zero Available Credits (≤ ₱5,000). Click an alert to jump to Internet Credits or Wallets. To switch roles: name at the top right → Logout → another demo account. Do not click Reset Demo Data unless you intend to wipe what you just walked.",
        italic=True,
        color=MUTED,
    )

    heading(doc, "3. Sub-Franchisee", 1)
    add_para(
        doc,
        "You sit between Admin and Franchisees. You buy Internet Credits from Admin, hold them as Available Credits, then release them to franchisees after they deposit cash (or via Direct Release). You also earn a share of every internet sale in your network.",
    )
    bullets(
        doc,
        [
            "Left menu you will use: Wallets, Internet Credits, Deposit Rates, Commission Settings, Transactions, Revenue, Reports.",
            "Dashboard and Franchisees appear in the menu but are not clickable in this demo.",
            "Your org in the sample data is Northern Mindanao Sub-Franchisee. Downlines: Franchisee A and Franchisee B.",
        ],
    )

    heading(doc, "3.1 Wallets — see inventory across your network", 2)
    add_para(
        doc,
        "After sign-in you land here. This is a monitor, not a place to move money. Credit loads still go through Internet Credits (cash + proof).",
    )
    shot(doc, "sub-wallets.png", "Figure 2. Sub-Franchisee Wallets — your inventory and every downline balance.")
    steps(
        doc,
        [
            "Click Wallets in the left menu (or stay here after login).",
            "Read Your Available Credits. That is stock you can still release to franchisees. After Reset Demo Data this is ₱135,000.00.",
            "Read Franchisee Credits and Retailer Credits. Those are balances already sitting with downlines, not your cash. Franchisees and retailers start at ₱0.",
            "Network Credits shows how many wallets you manage and how many are Low / Zero. Low means Available Credits ≤ ₱5,000.",
            "Optional: click the Low Balance chip above the table to show only wallets that need a top-up.",
            "If Your Available Credits is Low (≤ ₱5,000) or Zero, a warning/danger banner appears on this page (same idea as the bell). New Credits Request opens Internet Credits with the request sheet already open.",
            "Click the eye icon on a row to open that wallet’s details (Available Credits, parent, recent credit activity). There is no Minimum Balance card on this modal.",
            "If a franchisee is low, go to Internet Credits. You can wait for their request (your bell will ping) or use Direct Release after they have paid.",
        ],
    )

    heading(doc, "3.2 Internet Credits — buy from Admin", 2)
    add_para(
        doc,
        "This is how you restock. You deposit cash to Admin, attach proof, and wait for Admin to release credits into your Available Credits. Default buy rate is 60% (₱600 cash → 1,000 credits).",
    )
    shot(doc, "sub-credits-mine.png", "Figure 3. My Credits Request — your buys from Admin.")
    steps(
        doc,
        [
            "Click Internet Credits.",
            "Note the cards: Pending Deposits (cash from franchisees not yet released), Pending Credits, Credits Released, My Pending Requests, Available Credits.",
            "Click the My Credits Request tab.",
            "Click + New Credits Request at the top right. The New Credits Request sheet opens on the right.",
        ],
    )
    shot(doc, "sub-credits-new-request.png", "Figure 4. New Credits Request — deposit to Admin, suggested credits at 60%.")
    steps(
        doc,
        [
            "Enter Amount (PHP) — the cash you already sent to Admin. The form shows suggested credits using credits = deposit ÷ your 60% rate. Example: ₱30,000 → ₱50,000 credits.",
            "Optionally upload proof of payment (SVG, PNG, JPG, or PDF, max 5MB). Add notes if you have a bank reference. Proof is optional in the demo.",
            "Click Continue, then confirm. The row appears as Pending on My Credits Request. Type is Request.",
            "While it is still Pending, you can Update (amount, notes, proof) or Delete. Update reapplies the live deposit rate, not the original snapshot.",
            "When Admin releases, status becomes Released and Available Credits goes up. Open Released / History to see completed loads. There is no separate Transfer History tab.",
        ],
    )
    callout(
        doc,
        "What happens to cash here",
        "You pay cash to Admin. Until Admin releases, that cash is a pending deposit — not yet inventory. After release, you hold credits (an asset you can sell). You have not earned Internet Credits earnings yet; that happens only when a franchisee buys from you.",
    )

    heading(doc, "3.3 Internet Credits — sell / release to franchisees", 2)
    add_para(
        doc,
        "Franchisees deposit cash to you and submit a request. You check proof, then release credits from your Available Credits. Default sell rate is 70% (₱700 cash → 1,000 credits). That 10-point gap vs your 60% buy rate is your Internet Credits earnings.",
    )
    shot(doc, "sub-credits-downlines.png", "Figure 5. Downlines Credits Request — pending franchisee deposits.")
    steps(
        doc,
        [
            "Stay on Internet Credits. When a franchisee submits a request, the bell shows “{name} requested Internet Credits.” Click it to open this page.",
            "Click Downlines Credits Request. This tab lists pending requests only. Direct Release never queues here.",
            "Find a row with Status Pending. Deposit (PHP) is the cash they claim they sent. Credits may show ₱0.00 until you release.",
            "Click Review. Check name, deposited amount, deposit rate, suggested credits, and Available Credits After Release. If the after-release figure is red, buy from Admin first.",
            "Open Proof of Payment and match it to the deposited amount.",
            "To refuse, click Reject, enter a reason, and confirm. No credits move.",
            "To accept, click Approve & Release. Enter Payment reference ID. Confirm or edit Credits to release (pre-filled from deposit ÷ rate). Click Approve & Release.",
            "Credits leave your Available Credits and land in the franchisee’s wallet. Status becomes Released. Type stays Request.",
        ],
    )
    heading(doc, "Direct Release (skip the pending queue)", 3)
    add_para(
        doc,
        "Use this when the franchisee already paid and you do not need them to file a request first. Same economics: deposit ÷ rate, optional proof, payment reference required. The sheet is titled Direct Credits Release.",
    )
    steps(
        doc,
        [
            "On Internet Credits, click Direct Release (top right). The Direct Credits Release sheet opens on the right.",
        ],
    )
    shot(
        doc,
        "sub-credits-direct-release.png",
        "Figure 6. Direct Credits Release — pick a franchisee, enter cash deposited, see suggested credits.",
    )
    steps(
        doc,
        [
            "Pick the franchisee. Enter cash deposited. Confirm suggested credits (deposit ÷ that franchisee’s rate). Attach proof if you have it. Add a note if useful.",
            "Click Continue, then enter Payment reference ID and confirm. Credits move immediately. The row appears on Released / History and on the franchisee’s My Credits Request with Type = Direct Release.",
        ],
    )
    shot(doc, "sub-credits-released.png", "Figure 7. Released / History — completed loads in and out, including Direct Release.")
    add_para(
        doc,
        "Open Released / History to audit past releases (requests and Direct Releases). View is read-only. Reverse after release is not available in this version.",
    )
    add_para(doc, "Internet Credits earnings check on 1,000 credits sold to a franchisee at 70%:")
    bullets(
        doc,
        [
            "Cash in from franchisee: ₱700.",
            "Your cost (what you paid Admin at 60%): ₱600.",
            "Internet Credits earnings: ₱100. Same as 1,000 × (70% − 60%). This is inventory markup, not sale commission.",
        ],
    )

    heading(doc, "3.4 Deposit Rates — price the credit load", 2)
    add_para(
        doc,
        "This rate card is only for Internet Credits loads to franchisees. It does not change sale commissions. Demo starts with hop defaults only (no custom overrides).",
    )
    shot(doc, "sub-deposit-rates.png", "Figure 8. Deposit Rates — Sub-Franchisee → Franchisee default 70%.")
    steps(
        doc,
        [
            "Click Deposit Rates.",
            "The big card is your sell-hop default: Sub-Franchisee → Franchisee at 70%.",
            "Your downlines lists each franchisee. Source Default means they use 70%. Custom means you overrode that one account.",
            "To change one franchisee, click the pencil. Enter the new rate and a reason. Save. If you type the default 70% again, the override is cleared.",
            "New requests after the change use the new rate. If the requester Updates a still-pending request, the live rate is reapplied.",
        ],
    )

    heading(doc, "3.5 Commission Settings — split each internet sale", 2)
    add_para(
        doc,
        "This is the other pricing surface. Percentages apply to sales (customer payment), not to credit loads. For each retailer: Retailer %, Franchisee %, Your Share %, Platform Fee % (Admin-set in the demo). Total must be 100% of sales. These % are stamped on each sale — changing a row later does not rewrite completed sales.",
    )
    shot(doc, "sub-commission-settings.png", "Figure 9. Commission Settings — per-retailer sale splits.")
    steps(
        doc,
        [
            "Click Commission Settings.",
            "Optional: filter by Franchisee, Retailer, or Status, then Apply.",
            "Read a row. Demo seed: Retailer A, Retailer B, and Retailer C use 10% / 20% / 30% / 40% of sales.",
            "To change an existing row, click the pencil. You can edit Retailer %, Franchisee %, and Your Share. Platform fee is set by Admin and is read-only here. The four values must total 100%. Set the effective date and save.",
            "To add a retailer that has no row yet, click + Add Commission Settings, pick the retailer, enter the shares, and save. Saving Active inactivates the previous Active row for that retailer.",
            "These % apply to the customer payment of each completed internet sale, not to credit loads.",
        ],
    )

    heading(doc, "3.6 Transactions — see sales in your network", 2)
    add_para(
        doc,
        "This ledger is completed internet sales only. It is not the credit-load list. The demo starts empty — a retailer must Record demo sale before rows appear.",
    )
    shot(doc, "sub-transactions.png", "Figure 10. Sub-Franchisee Transactions Ledger.")
    steps(
        doc,
        [
            "Click Transactions.",
            "Optional: set Date Range, pick a Retailer, or search by name / ID, then Apply Filters.",
            "Read a row left to right: Customer Payment → Credits Consumed (red) → Your Share (your commission).",
            "After a ₱1,000 sale on Retailer A: customer ₱1,000, credits −₱1,000, your share ₱300.00 (30% of sales).",
            "Click the eye icon to open Transaction Details. Confirm product, parties, and Commission Distribution. Close when done.",
            "Optional: Export CSV for a spreadsheet.",
        ],
    )

    heading(doc, "3.7 Revenue — Internet Credits earnings plus sales commission", 2)
    add_para(
        doc,
        "This page answers “what did I earn?” It adds the two streams for the selected period.",
    )
    shot(doc, "sub-revenue.png", "Figure 11. Sub-Franchisee Revenue — Internet Credits earnings + Sales Commission.")
    steps(
        doc,
        [
            "Click Revenue.",
            "Read the three cards: Internet Credits earnings (spread on credits you released), Sales Commission (your % of sales), Total earnings (the sum).",
            "Change Date Range if needed (This Month is the default) and click Apply.",
            "Internet Credits table: each released downline load, with Cash in, Credits, and Earnings. Click the ⓘ next to an amount to see the formula (cash − cost, or credits × rate gap).",
            "Sales Commission table: each sale, with Sales, Your share %, and Your commission. Click the eye to open the same transaction detail.",
        ],
    )
    callout(
        doc,
        "How to read an Internet Credits earnings row",
        "If Franchisee A deposited ₱14,000 and received ₱20,000 credits at 70%, cash-in is ₱14,000. Your cost for those credits is ₱20,000 × your 60% buy rate = ₱12,000. Earnings = ₱2,000. That is inventory margin, not customer-sale commission.",
    )

    heading(doc, "3.8 Reports — period snapshot and CSV exports", 2)
    shot(doc, "sub-reports.png", "Figure 12. Sub-Franchisee Reports.")
    steps(
        doc,
        [
            "Click Reports.",
            "Set Date Range, Franchisee, and/or Retailer. Cards refresh for that period.",
            "Card order: Internet Credits earnings → Sales Commission → Total earnings → Sales Volume.",
            "Click the Internet Credits earnings card to open Revenue with the same period, scrolled to the line-item table (each release + ⓘ).",
            "Internet Credits by downline is a rollup of those same entries (downline, cash in, credits, earnings) — not a copy of the Revenue table.",
            "Franchisee Commissions / Retailer Commissions show Sales Volume and each party’s share of sales.",
            "Revenue Sharing Sub-Franchisee is the client-sheet table: Sales, Sub / Franchisee / Retailer shares, and Total Revenue (Sub + Franchisee + Retailer; platform fee is separate).",
            "Internet Retailer Balance Report: pick one retailer. Credits loaded versus credits consumed on sales, with running wallet balance. Franchisee column is the parent franchisee.",
            "Downline Credit Statement: pick a franchisee or retailer and read + / − movements like a bank statement. Download becomes available after you select a downline (CSV of + / − and running balance).",
            "Scroll to Export Reports at the bottom of the page for the CSV cards.",
        ],
    )
    shot(doc, "sub-reports-export.png", "Figure 13. Sub-Franchisee Export Reports — CSV downloads for the selected period.")
    add_para(
        doc,
        "Each card uses the same Date Range / Franchisee / Retailer filters as the rest of Reports. The count is how many rows will be in the file. Export CSV is disabled when the count is 0.",
    )
    add_table(
        doc,
        ["CSV", "What’s in the file"],
        [
            [
                "Transactions",
                "Each sale in your scope: customer payment, credits consumed, split %, and your share.",
            ],
            [
                "Sales Commission",
                "Each sale: sales, your share %, your commission, status.",
            ],
            [
                "Revenue Sharing Sub-Franchisee",
                "Sales and Sub / Franchisee / Retailer shares. Total Revenue excludes platform fee.",
            ],
            [
                "Internet Credits earnings",
                "Each released downline load: cash in, credits, earnings, buy rate, sell rate.",
            ],
            [
                "Internet Credits Requests",
                "Requests you can see: deposit, suggested/released credits, rate, payment ref, status (includes Direct Release).",
            ],
            [
                "Internet Credits Releases",
                "Credit release ledger: deposit, rate, and payment reference.",
            ],
        ],
    )
    steps(
        doc,
        [
            "Set filters first (This Month is the default), then scroll to Export Reports.",
            "Click Export CSV on the card you need. The browser downloads a file named like esarisari-transactions-report-subfranchisee.csv.",
            "Open it in Excel, Google Sheets, or Numbers. Amounts are unformatted numbers so you can sum them.",
            "On Franchisee Commissions / Retailer Commissions tables, the download icon on a row exports that party’s sale lines only.",
        ],
    )
    callout(
        doc,
        "Exports follow the filters",
        "If you set Date Range to Last 7 Days and Franchisee to Franchisee A, every Export CSV card (and the statement Download) uses that slice. Change filters, then export again for a new file.",
        fill="EEF2FF",
    )

    heading(doc, "4. Franchisee", 1)
    add_para(
        doc,
        "You sit between the Sub-Franchisee and retailers. You buy credits from your Sub-Franchisee (demo buy rate 70%), hold Available Credits, then release to retailers (demo sell rate 80%). You also earn a share of each retailer sale. Sign in as Franchisee A for this walkthrough.",
    )
    bullets(
        doc,
        [
            "Left menu: Wallets, Internet Credits, Deposit Rates, Transactions, Revenue, Reports.",
            "You do not have Commission Settings — sale splits are set by the Sub-Franchisee / Admin.",
            "Dashboard and Retailers appear but are not clickable in this demo.",
            "Franchisee A’s retailers are Retailer A and Retailer B. You start at ₱0 Available Credits.",
        ],
    )

    heading(doc, "4.1 Wallets — you and your retailers", 2)
    shot(doc, "fran-wallets.png", "Figure 14. Franchisee Wallet Management.")
    steps(
        doc,
        [
            "Sign in as Franchisee A, or click Wallets.",
            "Your Available Credits is stock you can still release to retailers. After a fresh demo this is ₱0 — restock via Internet Credits first. A Zero (red) or Low (amber) banner on this page also appears when your own inventory is at or below ₱5,000; New Credits Request opens the request sheet to your Sub-Franchisee.",
            "Retailer Credits is inventory already with your retailers. Managed Credits counts wallets and flags low / zero (≤ ₱5,000).",
            "Click Low Balance to filter retailers who need a top-up.",
            "Click the eye on a retailer to see Available Credits, parent (you), and recent credits received. There is no Minimum Balance card.",
        ],
    )
    shot(doc, "fran-wallet-details.png", "Figure 15. Wallet details — Available Credits and recent activity.")
    add_para(
        doc,
        "A Low Balance badge means Available Credits is at or below ₱5,000. It does not move money by itself. The retailer must deposit cash and you release from Internet Credits (or you Direct Release after they paid).",
    )

    heading(doc, "4.2 Internet Credits — buy from your Sub-Franchisee", 2)
    add_para(
        doc,
        "Same buy pattern as the Sub-Franchisee, one hop down. You deposit cash to Northern Mindanao Sub-Franchisee and wait for them to release. Default: ₱700 cash → 1,000 credits. Their bell will show your request.",
    )
    shot(doc, "fran-credits-mine.png", "Figure 16. Franchisee My Credits Request — buys from the Sub-Franchisee.")
    steps(
        doc,
        [
            "Click Internet Credits.",
            "Cards: Pending Deposits (retailer cash waiting on you), Pending Credits, Credits Released, My Pending Requests, Available Credits.",
            "Click My Credits Request.",
            "Click + New Credits Request. The New Credits Request sheet opens on the right.",
        ],
    )
    shot(
        doc,
        "fran-credits-new-request.png",
        "Figure 17. New Credits Request — deposit to the Sub-Franchisee, suggested credits at 70%.",
    )
    steps(
        doc,
        [
            "Enter the cash you deposited to the Sub-Franchisee. Suggested credits = deposit ÷ your buy rate (70% unless overridden).",
            "Optionally attach proof and notes → Continue → confirm.",
            "While Pending, you may Update or Delete. When Released (or when you receive a Direct Release), your Available Credits increases. Type shows Request or Direct Release.",
        ],
    )

    heading(doc, "4.3 Internet Credits — sell / release to retailers", 2)
    add_para(
        doc,
        "Retailers deposit cash to you. You release from Available Credits at the per-retailer deposit rate (default 80%). Your Internet Credits earnings are sell-rate cash minus your 70% cost.",
    )
    shot(doc, "fran-credits-retailers.png", "Figure 18. Retailers Credits Request — pending retailer deposits.")
    steps(
        doc,
        [
            "The bell alerts you when a retailer requests credits. Click it, or open Internet Credits → Retailers Credits Request.",
            "Open a Pending row with Review. Check Deposited Amount, deposit rate, suggested credits, and Available Credits After Release.",
            "Match Proof of Payment. Reject with a reason if the proof is wrong.",
            "Click Approve & Release. Enter Payment reference ID. Confirm Credits to release. Save.",
            "Or click Direct Release to load a retailer immediately after they paid, without a pending row. The Direct Credits Release sheet opens on the right.",
        ],
    )
    shot(
        doc,
        "fran-credits-direct-release.png",
        "Figure 19. Direct Credits Release — pick a retailer, enter cash deposited, see suggested credits at 80%.",
    )
    steps(
        doc,
        [
            "Pick the retailer, enter cash deposited, confirm suggested credits, then Continue → Payment reference ID → confirm.",
            "Credits leave your inventory and appear on the retailer’s Wallet / Internet Credits.",
        ],
    )
    shot(doc, "fran-credits-released.png", "Figure 20. Franchisee Released / History.")
    add_para(doc, "Internet Credits earnings check on 1,000 credits sold to a retailer at 80%:")
    bullets(
        doc,
        [
            "Cash in from retailer: ₱800.",
            "Your cost (paid to Sub-Franchisee at 70%): ₱700.",
            "Internet Credits earnings: ₱100. If you sold at 70%, spread would be ₱0 — that is why Deposit Rates matter.",
        ],
    )

    heading(doc, "4.4 Deposit Rates — what retailers pay you", 2)
    shot(doc, "fran-deposit-rates.png", "Figure 21. Franchisee Deposit Rates — default 80%.")
    steps(
        doc,
        [
            "Click Deposit Rates.",
            "Default hop is Franchisee → Retailer at 80%. Demo starts with Default on every retailer.",
            "Edit a retailer with the pencil to set a custom rate (and a reason). Resetting to 80% removes the custom flag.",
            "Internet Credits earnings on Revenue use this rate card. Sale commissions still use the distribution % on each sale.",
        ],
    )

    heading(doc, "4.5 Transactions — retailer sales and your share", 2)
    shot(doc, "fran-transactions.png", "Figure 22. Franchisee Transactions Ledger.")
    steps(
        doc,
        [
            "Click Transactions. The list is empty until a retailer records a demo sale.",
            "Filter by Date Range or Retailer, search, then Apply Filters.",
            "Your Share is your commission from that sale’s customer payment (20% of sales on the demo seed), not the full ₱1,000.",
            "Click the eye on a row for the full split.",
        ],
    )
    shot(doc, "fran-transaction-details.png", "Figure 23. Transaction Details — commission split of sales across the chain.")
    add_para(
        doc,
        "On a ₱1,000 Retailer A sale you should see: customer ₱1,000, credits −₱1,000. Your commission is ₱200.00 (20% of sales). Retailer ₱100.00, Sub-Franchisee ₱300.00, platform ₱400.00. Optional: Download Receipt, then Close.",
    )

    heading(doc, "4.6 Revenue and Reports", 2)
    shot(doc, "fran-revenue.png", "Figure 24. Franchisee Revenue — loads to retailers plus sales commission.")
    steps(
        doc,
        [
            "Click Revenue. Internet Credits earnings = spread on credits you released to retailers. Sales Commission = your % of sales. Total earnings = both.",
            "Internet Credits table lists each retailer load (Cash in, Credits, Earnings) with ⓘ on the amount. Sales Commission lists each sale (Sales, Your share %, Your commission).",
            "Click Reports. Same period cards plus Sales Volume. Click the Internet Credits earnings card to open Revenue for that period.",
            "Internet Credits by downline rolls up cash in, credits, and earnings per retailer. Retailer Commissions uses Sales Volume and each party’s share of sales.",
            "Internet Retailer Balance Report: pick one of your retailers for credits loaded versus sales, with running wallet balance.",
            "Scroll to Export Reports for CSV downloads (same filters as the cards above).",
        ],
    )
    shot(doc, "fran-reports.png", "Figure 25. Franchisee Reports.")
    shot(doc, "fran-reports-export.png", "Figure 26. Franchisee Export Reports — same CSV cards, scoped to your retailers.")
    add_para(
        doc,
        "Same files as the Sub-Franchisee except Revenue Sharing Sub-Franchisee (that table is Admin / Sub only): Transactions, Sales Commission, Internet Credits earnings, Requests, and Releases. Counts and rows are only your network (your restock from the Sub plus loads to Retailer A / B, and sales under you). Click Export CSV; the filename includes franchisee.",
    )
    callout(
        doc,
        "Pending loads are not earnings",
        "Internet Credits earnings only count released downline loads (request or Direct Release). A pending My Credits Request to the Sub-Franchisee is not yet earned. Your own restock is a purchase of inventory, not revenue.",
        fill="FFF7ED",
    )

    heading(doc, "5. Retailer", 1)
    add_para(
        doc,
        "You are the end of the credit chain and the start of customer sales. You only buy credits (you have no downline to approve). You deposit cash to your Franchisee, receive Available Credits, then burn those credits when you record a demo sale. You keep your % of each sale’s customer payment — 10% on the demo seed (₱100 on ₱1,000) — not the full payment. Use Retailer A under Franchisee A.",
    )
    bullets(
        doc,
        [
            "Left menu: Wallet, Internet Credits, Transactions, Revenue, Reports.",
            "No Deposit Rates, Commission Settings, or Direct Release — those sit above you.",
            "You start at ₱0 Available Credits. Demo buy rate is 80% unless your franchisee set a custom rate.",
        ],
    )

    heading(doc, "5.1 Wallet — inventory for sales", 2)
    shot(doc, "ret-wallet.png", "Figure 27. Retailer Wallet — Available Credits and upline.")
    steps(
        doc,
        [
            "Sign in as Retailer A. You land on Wallet. The page tells you to request more via Internet Credits from your franchisee.",
            "Available Credits is what you can still sell. Status Low / Zero (≤ ₱5,000, or ₱0) means restock soon. After a fresh demo you are at ₱0.",
            "Franchisee shows who you buy from (Franchisee A).",
            "Recent Credits Activity lists credits received after the franchisee released your request or Direct Release.",
        ],
    )
    shot(
        doc,
        "ret-wallet-low-balance.png",
        "Figure 28. Retailer B Wallet — Low Balance banner with New Credits Request.",
    )
    add_para(
        doc,
        "Sign in as Retailer B to see the on-page alert. In this demo Retailer B has ₱2,500 Available Credits, which is ₱5,000 or less, so Wallet shows an amber Low Balance banner in addition to the bell. Zero Balance uses a red danger banner instead.",
    )
    steps(
        doc,
        [
            "Read the banner: it shows the live Available Credits and that we alert when the balance is ₱5,000 or less.",
            "Click New Credits Request on the banner. That opens Internet Credits with the New Credits Request sheet already open — you do not have to hunt for + New Credits Request.",
            "Complete the request (deposit, optional proof) the same way as section 5.2. The franchisee’s bell will show it.",
        ],
    )

    heading(doc, "5.2 Internet Credits — request stock from the Franchisee", 2)
    add_para(
        doc,
        "There is no downline tab and no Direct Release. You only request. Credits = deposit ÷ your rate. At 80%, ₱800 cash → 1,000 credits. Your franchisee’s bell will show the request.",
    )
    shot(doc, "ret-credits-mine.png", "Figure 29. Retailer My Credits Request.")
    steps(
        doc,
        [
            "Click Internet Credits.",
            "My Pending Requests is cash you sent that is not released yet. Available Credits is live inventory.",
            "Click + New Credits Request. The New Credits Request sheet opens on the right.",
        ],
    )
    shot(
        doc,
        "ret-credits-new-request.png",
        "Figure 30. New Credits Request — deposit to Franchisee A, suggested credits at 80%.",
    )
    steps(
        doc,
        [
            "Enter Amount (PHP) you already sent to Franchisee A. The calculator shows suggested credits at your deposit rate.",
            "Optionally upload proof and notes → Continue → confirm.",
            "While Pending, you may Update or Delete. Update reapplies the live rate.",
            "When the franchisee releases (or Direct Releases), status becomes Released and Wallet Available Credits goes up. Type shows Request or Direct Release.",
            "Click Released / History to see completed loads.",
        ],
    )
    shot(doc, "ret-credits-released.png", "Figure 31. Retailer Released / History.")
    callout(
        doc,
        "Your cash on a load is a purchase of inventory",
        "₱1,600 cash for ₱2,000 credits is not “profit.” You spent cash to buy stock. Profit from customers shows later as Your Commission on each sale, after credits are consumed.",
        fill="ECFDF5",
    )

    heading(doc, "5.3 Transactions — record and read each customer sale", 2)
    add_para(
        doc,
        "This is the only role that can create a sale in the demo. Record demo sale burns Available Credits and stamps commission shares for the whole chain.",
    )
    shot(doc, "ret-transactions.png", "Figure 32. Retailer Transactions Ledger.")
    steps(
        doc,
        [
            "Click Transactions.",
            "Click Record demo sale. Pick a product/service (or type one) and enter the customer payment. The preview shows credits that will be consumed (100% of payment) and remaining Available Credits.",
            "If credits required are more than your Available Credits, restock first. Submit to post a completed sale.",
            "Read Customer Payment, Credits Consumed, then Your Share (your commission).",
            "Click the eye. Transaction Details shows product, the chain, and Commission Distribution with You highlighted.",
            "Optional: Download Receipt, then Close. Export CSV if you need a file.",
        ],
    )
    shot(doc, "ret-transaction-details.png", "Figure 33. Retailer view of a sale — you earn your % of sales.")
    add_para(doc, "Same ₱1,000 sale, retailer reading:")
    bullets(
        doc,
        [
            "Customer paid ₱1,000. That cash is collected at the store. It is not all yours.",
            "₱1,000 of credits left your Wallet. That is the inventory cost of the sale.",
            "Your 10% of sales on Retailer A = ₱100.00. The rest of the ₱1,000 is franchisee, sub-franchisee, and platform shares.",
        ],
    )

    heading(doc, "5.4 Revenue — your commission only", 2)
    add_para(
        doc,
        "Retailers do not have an Internet Credits earnings table. You do not earn a credit spread — you buy at the franchisee’s rate and sell to customers.",
    )
    shot(doc, "ret-revenue.png", "Figure 34. Retailer Revenue — commission, sales volume, credits consumed.")
    steps(
        doc,
        [
            "Click Revenue.",
            "Your Commission is the sum of your sale shares for the period.",
            "Sales volume is customer payments. Credits consumed is inventory burned (equal to the payment). Your commission is your % of sales, not the full ₱1,000.",
            "Change Date Range / search and Apply as needed.",
            "Sales Commission table lists each sale: Sales, Your share %, Your commission, Status. Click the eye for details.",
        ],
    )

    heading(doc, "5.5 Reports", 2)
    shot(doc, "ret-reports.png", "Figure 35. Retailer Reports.")
    steps(
        doc,
        [
            "Click Reports.",
            "You see Sales Commission and Sales Volume only — no Internet Credits earnings card and no by-downline rollup (you have no load-spread stream).",
            "Internet Retailer Balance Report is your own wallet: Initial Wallet / Credit Received versus Sales, with running Wallet Balance.",
            "Downline Credit Statement on this role is your own statement — opening balance, + releases, − sales, ending balance. Download exports that statement as CSV.",
            "Scroll to Export Reports. You do not get an Internet Credits earnings CSV (you have no load-spread stream).",
        ],
    )
    shot(doc, "ret-reports-export.png", "Figure 36. Retailer Export Reports — Transactions, Sales Commission, Requests, and Releases.")
    add_para(
        doc,
        "Four CSVs: Transactions (your sales), Sales Commission (your % of each sale), Internet Credits Requests (your buys from the franchisee), and Internet Credits Releases (credits that landed in your Wallet). Same rule as other roles: filters first, then Export CSV. Filename includes retailer.",
    )

    heading(doc, "6. End-to-end flow (all three roles)", 1)
    add_para(doc, "Use this when training a new operator. Reset Demo Data first so everyone starts empty except the Sub-Franchisee’s ₱135,000 stock.")

    heading(doc, "A. Restock credits (cash up the chain, credits down)", 2)
    steps(
        doc,
        [
            "Retailer A deposits cash to Franchisee A and submits Internet Credits → + New Credits Request with proof. Franchisee A’s bell shows the request.",
            "Franchisee A cannot release yet (₱0 stock). Franchisee A submits My Credits Request to the Sub-Franchisee (or the Sub uses Direct Release after cash is confirmed).",
            "Sub-Franchisee opens the bell or Downlines Credits Request → Review → Approve & Release (payment ref + credits). Sub Available Credits down; Franchisee A up.",
            "Franchisee A opens Retailers Credits Request (or Direct Release) → Approve & Release. Franchisee Available Credits down; Retailer Wallet up.",
            "When the Sub-Franchisee is low, they submit My Credits Request to Admin (not covered in this guide) and wait for release.",
        ],
    )

    heading(doc, "B. One internet sale (customer cash, then the split)", 2)
    steps(
        doc,
        [
            "Retailer A opens Transactions → Record demo sale (example Smart load ₱1,000). Wallet burns ₱1,000 credits if stock is enough.",
            "Commission Settings % of the ₱1,000 sales are stamped: retailer 10 / franchisee 20 / sub-franchisee 30 / platform 40.",
            "Each role opens Transactions → eye icon to see their line. Revenue / Reports add those lines for the month. Sub and Fran also see Internet Credits earnings from the restock in step A.",
        ],
    )

    heading(doc, "7. Field meanings (quick reference)", 1)
    add_table(
        doc,
        ["You see", "It means", "It is not"],
        [
            ["Available Credits", "Credit inventory you can sell or release", "Cash in your bank"],
            ["Pending Deposits", "Cash claimed, not yet released", "Already earned revenue"],
            ["Direct Release", "Same load as a request, no pending queue", "A free transfer without cash / proof / payment ref"],
            ["Internet Credits earnings", "Spread (or Admin cash) on credits you sold to a downline", "Customer-sale commission"],
            ["Sales Commission / Your Share", "Your % of the customer payment (Sales)", "The full ₱1,000"],
            ["Low Balance", "Available Credits ≤ ₱5,000", "A second wallet or a cash account"],
            ["Credits Consumed", "Inventory burned on a sale", "Your commission"],
        ],
    )

    heading(doc, "8. Common mistakes", 1)
    bullets(
        doc,
        [
            "Releasing credits without matching proof of payment — cash and inventory will not reconcile.",
            "Releasing when Available Credits After Release is negative — restock from your upline first (Franchisee A starts at ₱0).",
            "Treating the full customer payment as retailer profit — on a ₱1,000 Retailer A sale the retailer earns ₱100 (10% of sales).",
            "Changing Deposit Rates and expecting sale commissions to move — they are separate.",
            "Calling Internet Credits earnings “Load earnings” or mixing it with Sales Commission on Reports.",
            "Expecting Reverse after a mistaken release — that action is not in the UI.",
            "Using Reset Demo Data during training — it wipes balances, requests, and sales back to the empty demo.",
        ],
    )

    heading(doc, "9. Logout", 1)
    steps(
        doc,
        [
            "Click your name at the top right.",
            "Click Logout.",
            "Sign in with the next demo account if you are walking another role.",
        ],
    )

    add_para(
        doc,
        "This guide matches the eSariSari demo build as of 2 September 2026. Admin screens are omitted on purpose. Screenshots are layout references; live numbers depend on the walkthrough you just ran.",
        italic=True,
        color=MUTED,
        space_before=12,
    )

    doc.save(OUTPUT)
    OUTPUT_HTML.write_text(HTML.render(), encoding="utf-8")
    print(f"Wrote {OUTPUT}")
    print(f"Wrote {OUTPUT_HTML}")

    import sys

    sys.path.insert(0, str(ROOT))
    from _generate_computations_xlsx import build as build_xlsx

    build_xlsx()


if __name__ == "__main__":
    build()
