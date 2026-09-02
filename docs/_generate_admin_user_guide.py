#!/usr/bin/env python3
"""Generate docs/admin-user-guide.docx and docs/admin-user-guide.html.

Does not touch the Internet Credits user guide (docs/user-guide.*).
"""

from html import escape
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "admin-user-guide-assets"
OUTPUT = ROOT / "admin-user-guide.docx"
OUTPUT_HTML = ROOT / "admin-user-guide.html"

GUIDE_TITLE = "eSariSari Admin User Guide"
GUIDE_VERSION = "1.1"
GUIDE_DATE = "2 September 2026"
GUIDE_EDITION = f"Version {GUIDE_VERSION}  ·  {GUIDE_DATE}  ·  Demo environment"

NAVY = RGBColor(0x0F, 0x17, 0x2A)
BLUE = RGBColor(0x1D, 0x4E, 0xD8)
MUTED = RGBColor(0x47, 0x55, 0x69)
AMBER = RGBColor(0xB4, 0x53, 0x09)


class HtmlSink:
    def __init__(self):
        self.body = []
        self.toc = []
        self._slugs = set()

    def slug(self, text):
        base = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-") or "section"
        slug = base
        n = 2
        while slug in self._slugs:
            slug = f"{base}-{n}"
            n += 1
        self._slugs.add(slug)
        return slug

    def add(self, html):
        self.body.append(html)

    def heading(self, text, level):
        slug = self.slug(text)
        if level <= 2:
            self.toc.append((level, slug, text))
        self.add(f'<h{level} id="{slug}">{escape(text)}</h{level}>')

    def render(self):
        toc_items = []
        for level, slug, text in self.toc:
            cls = "toc-h1" if level == 1 else "toc-h2"
            toc_items.append(
                f'<a class="{cls}" href="#{slug}">{escape(text)}</a>'
            )
        return HTML_PAGE.format(
            title=escape(f"{GUIDE_TITLE} · Version {GUIDE_VERSION}"),
            toc="\n".join(toc_items),
            body="\n".join(self.body),
        )


HTML = HtmlSink()

HTML_PAGE = """<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>{title}</title>
  <style>
    :root {{
      --navy: #0f172a;
      --blue: #1d4ed8;
      --muted: #475569;
      --line: #e2e8f0;
      --bg: #f8fafc;
    }}
    * {{ box-sizing: border-box; }}
    html {{ scroll-behavior: smooth; }}
    body {{
      margin: 0;
      font-family: Calibri, "Segoe UI", system-ui, sans-serif;
      color: var(--navy);
      background: var(--bg);
      line-height: 1.5;
    }}
    .wrap {{
      display: grid;
      grid-template-columns: 16rem minmax(0, 48rem);
      gap: 2rem;
      max-width: 72rem;
      margin: 0 auto;
      padding: 2rem 1.25rem 4rem;
    }}
    nav {{
      position: sticky;
      top: 1rem;
      align-self: start;
      max-height: calc(100vh - 2rem);
      overflow: auto;
      padding-right: 0.5rem;
    }}
    nav .label {{
      font-size: 0.7rem;
      font-weight: 700;
      letter-spacing: 0.08em;
      text-transform: uppercase;
      color: var(--muted);
      margin-bottom: 0.75rem;
    }}
    nav a {{
      display: block;
      color: var(--muted);
      text-decoration: none;
      font-size: 0.85rem;
      padding: 0.2rem 0;
    }}
    nav a:hover {{ color: var(--blue); }}
    nav .toc-h1 {{ font-weight: 700; color: var(--navy); margin-top: 0.6rem; }}
    nav .toc-h2 {{ padding-left: 0.75rem; font-size: 0.8rem; }}
    article {{
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 1rem;
      padding: 2rem 1.75rem 3rem;
    }}
    .kicker {{ color: var(--blue); font-weight: 700; margin: 0; }}
    .lede {{ color: var(--muted); margin: 0 0 0.5rem; }}
    h1.cover {{ font-size: 2.25rem; margin: 0.25rem 0 0.5rem; }}
    h1 {{ font-size: 1.55rem; margin: 2rem 0 0.6rem; }}
    h2 {{ font-size: 1.2rem; margin: 1.5rem 0 0.5rem; }}
    h3 {{ font-size: 1.05rem; margin: 1.25rem 0 0.4rem; }}
    p {{ margin: 0 0 0.85rem; }}
    p.muted {{ color: var(--muted); font-style: italic; }}
    ol, ul {{ margin: 0 0 1rem; padding-left: 1.25rem; }}
    li {{ margin: 0.25rem 0; }}
    ol li::marker {{ color: var(--blue); font-weight: 700; }}
    table {{
      width: 100%;
      border-collapse: collapse;
      font-size: 0.9rem;
      margin: 0 0 1.1rem;
    }}
    th, td {{
      border: 1px solid var(--line);
      padding: 0.45rem 0.55rem;
      text-align: left;
      vertical-align: top;
    }}
    th {{ background: var(--navy); color: #fff; font-size: 0.78rem; }}
    tr:nth-child(even) td {{ background: #f8fafc; }}
    .callout {{
      border: 1px solid #c7d2fe;
      border-radius: 0.75rem;
      padding: 0.85rem 1rem;
      margin: 0 0 1rem;
    }}
    .callout strong {{ color: var(--blue); display: block; margin-bottom: 0.25rem; }}
    figure {{ margin: 1rem 0 1.25rem; text-align: center; }}
    figure img {{
      max-width: 100%;
      height: auto;
      border: 1px solid var(--line);
      border-radius: 0.5rem;
    }}
    figcaption {{ color: var(--muted); font-size: 0.85rem; font-style: italic; margin-top: 0.4rem; }}
    .missing {{ color: #b45309; font-style: italic; }}
    .download {{ font-size: 0.9rem; margin: 0 0 1.25rem; }}
    .download a {{ color: var(--blue); }}
    @media (max-width: 860px) {{
      .wrap {{ grid-template-columns: 1fr; }}
      nav {{ position: static; max-height: none; }}
    }}
    @media print {{
      nav {{ display: none; }}
      .wrap {{ display: block; max-width: none; padding: 0; }}
      article {{ border: 0; }}
    }}
  </style>
</head>
<body>
  <div class="wrap">
    <nav>
      <div class="label">Contents</div>
      {toc}
    </nav>
    <article>
      {body}
    </article>
  </div>
</body>
</html>
"""


def set_run(run, *, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="CBD5E1"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_para(doc, text, *, size=11, bold=False, color=None, space_after=8, space_before=0, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color, italic=italic)
    if size >= 28:
        HTML.add(f'<h1 class="cover">{escape(text)}</h1>')
    elif size >= 14 and color == BLUE:
        weight = "700" if bold else "600"
        HTML.add(f'<p class="kicker" style="font-weight:{weight}">{escape(text)}</p>')
    elif size >= 11 and color == MUTED and not italic:
        HTML.add(f'<p class="lede">{escape(text)}</p>')
    else:
        cls = ' class="muted"' if italic or color == MUTED else ""
        HTML.add(f"<p{cls}>{escape(text)}</p>")
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(16 if level > 1 else 8)
    p.paragraph_format.space_after = Pt(8)
    HTML.heading(text, 3 if level >= 3 else level)
    return p


def shot(doc, filename, caption):
    path = ASSETS / filename
    if not path.exists():
        add_para(doc, f"[Screenshot missing: {filename}]", italic=True, color=AMBER)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.3))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    set_run(r, size=9, italic=True, color=MUTED)
    HTML.add(
        "<figure>"
        f'<img src="admin-user-guide-assets/{escape(filename)}" alt="{escape(caption)}" />'
        f"<figcaption>{escape(caption)}</figcaption>"
        "</figure>"
    )


def steps(doc, items):
    lis = []
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        n = p.add_run(f"{i}.  ")
        set_run(n, size=11, bold=True, color=BLUE)
        t = p.add_run(item)
        set_run(t, size=11, color=NAVY)
        lis.append(f"<li>{escape(item)}</li>")
    HTML.add("<ol>" + "".join(lis) + "</ol>")


def bullets(doc, items):
    lis = []
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(3)
        n = p.add_run("•  ")
        set_run(n, size=11, color=BLUE)
        t = p.add_run(item)
        set_run(t, size=11, color=NAVY)
        lis.append(f"<li>{escape(item)}</li>")
    HTML.add("<ul>" + "".join(lis) + "</ul>")


def callout(doc, title, body, fill="EEF2FF"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_borders(cell, "C7D2FE")
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(title)
    set_run(r1, size=10, bold=True, color=BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(body)
    set_run(r2, size=10, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    HTML.add(
        f'<div class="callout" style="background:#{fill.lower()}">'
        f"<strong>{escape(title)}</strong>{escape(body)}</div>"
    )


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "0F172A")
        set_cell_borders(cell, "0F172A")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run(r, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for ri, row in enumerate(rows):
        for ci, value in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            shade_cell(cell, "F8FAFC" if ri % 2 == 0 else "FFFFFF")
            set_cell_borders(cell, "E2E8F0")
            p = cell.paragraphs[0]
            r = p.add_run(str(value))
            set_run(r, size=9, color=NAVY)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)
    head = "".join(f"<th>{escape(h)}</th>" for h in headers)
    body = "".join(
        "<tr>" + "".join(f"<td>{escape(str(value))}</td>" for value in row) + "</tr>"
        for row in rows
    )
    HTML.add(f"<table><thead><tr>{head}</tr></thead><tbody>{body}</tbody></table>")


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    core = doc.core_properties
    core.title = GUIDE_TITLE
    core.subject = f"Platform Admin — Financials Dashboard and Client Onboarding · Version {GUIDE_VERSION}"
    core.author = "eSariSari"
    core.revision = 2
    core.version = GUIDE_VERSION
    core.comments = f"Version {GUIDE_VERSION} — {GUIDE_DATE}"

    add_para(doc, "eSariSari", size=14, bold=True, color=BLUE, space_after=0)
    add_para(doc, "Franchise financial operations platform", size=11, color=MUTED, space_after=18)
    add_para(doc, "Admin User Guide", size=28, bold=True, color=NAVY, space_after=4)
    add_para(
        doc,
        "Platform Admin  ·  Financials Dashboard  ·  Clients",
        size=14,
        color=BLUE,
        space_after=12,
    )
    add_para(
        doc,
        "A click-by-click walkthrough of the Admin franchise module: sign in, read the Financials Dashboard, confirm collections, open clients, and add a new client through the four-step onboarding wizard. Create Franchisee saves that client to this browser and shows it on the Clients list. Internet Credits, wallets, sales, and downline roles are covered in the separate User Guide — not in this document.",
        size=11,
        color=NAVY,
        space_after=10,
    )
    add_para(doc, GUIDE_EDITION, size=10, italic=True, color=MUTED)
    HTML.add(
        '<p class="download"><a href="admin-user-guide.docx">Download Admin User Guide Word copy</a> · '
        '<a href="user-guide.html">Open the Internet Credits User Guide</a></p>'
    )

    callout(
        doc,
        f"About this demo (version {GUIDE_VERSION})",
        "This is a browser-only demo. Data lives in localStorage. There is no backend. Reset Demo Data restores the seeded client portfolio, removes clients you registered, clears activation overlays, and clears onboarding drafts and collection marks. Screenshots show layout; peso amounts change as you confirm collections.",
        fill="EEF2FF",
    )

    heading(doc, "What's new in this version", 1)
    add_para(
        doc,
        f"Version {GUIDE_VERSION} matches the Admin module as of {GUIDE_DATE}.",
    )
    add_table(
        doc,
        ["Version", "Date", "What changed"],
        [
            [
                "1.1 (current)",
                "2 September 2026",
                "Confirm Collection on Transactions, Revenue, and Reports (same ledger as Dashboard and Client Details). Collection cash stays out of Total earnings.",
            ],
            [
                "1.0",
                "1 September 2026",
                "First Admin walkthrough: Financials Dashboard, Clients, onboarding wizard, Activate Client.",
            ],
        ],
    )

    heading(doc, "1. What this guide covers", 1)
    add_para(
        doc,
        "Use this guide when you are signed in as Platform Admin. The Admin home page is the Financials Dashboard. Clients is the franchise portfolio and the entry point for onboarding.",
    )
    add_table(
        doc,
        ["In this guide", "Not in this guide"],
        [
            [
                "Financials Dashboard KPIs and Confirm Collection",
                "Internet Credits loads, Direct Release, and deposit rates",
            ],
            [
                "Clients list, client details, and shared collection history",
                "Recording a retailer demo sale or reading Sales Commission",
            ],
            [
                "Add New Client — Client Info, Franchise Setup, Revenue Split, Review",
                "Sub-Franchisee, Franchisee, and Retailer day-to-day screens",
            ],
            [
                "Franchise collections on Transactions, Revenue, and Reports (Admin)",
                "Mixing collection cash into Internet Credits + sales Total earnings",
            ],
        ],
    )
    callout(
        doc,
        "Two separate user guides",
        "Open User Guide (or user-guide.html) for Internet Credits and the Sub / Fran / Retailer walkthrough. Open Admin User Guide (this file) for franchise setup, collections, and onboarding.",
        fill="ECFDF5",
    )

    heading(doc, "2. Sign in as Admin", 1)
    add_para(
        doc,
        "The Admin demo account is marked Restricted on the Sign In page. Clicking a demo account fills the email only — you still type the password.",
    )
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Email", "admin@esarisari.local"],
            ["Password", "abc12345678"],
            ["Other demo accounts", "password123"],
        ],
    )
    steps(
        doc,
        [
            "Open the app and stay on Sign In.",
            "Under Demo Accounts, click Admin (Restricted).",
            "Type abc12345678 in Password. Do not use password123 for Admin.",
            "Click Sign In.",
        ],
    )
    shot(doc, "admin-01-login.png", "Sign In with the Restricted Admin demo account selected.")
    add_para(
        doc,
        "After a successful sign-in you land on Financials Dashboard. The sidebar for Admin is Dashboard, Clients, Wallets, Internet Credits, Deposit Rates, Commission Settings, Transactions, Revenue, and Reports. Franchises and Organizations are not in the sidebar.",
    )

    heading(doc, "3. Financials Dashboard", 1)
    add_para(
        doc,
        "Dashboard is the Admin home. It tracks setup fees and monthly collections for the franchise portfolio — seeded demo clients plus any clients you register in this browser. Newly registered clients count under Pending Review. Open Franchise Setup in the page header jumps to onboarding step 1 (same wizard as Add New Client).",
    )
    shot(doc, "admin-02-dashboard.png", "Financials Dashboard — General KPIs, commitments by status, and collections.")

    heading(doc, "3.1 General KPIs", 2)
    add_table(
        doc,
        ["Card", "What it means"],
        [
            ["Active Upfront", "One-time setup due from Activated clients, plus how many of those upfront items are already collected."],
            ["Active Billable Monthly", "Fixed monthly fees due from Activated clients for the selected month. Cost-deduction-only and % of gross sales items are not in this billable total."],
            ["Activated Portfolio", "Count of Activated clients and their configured upfront / monthly commitments."],
            ["Coverage", "Territory and area count across the portfolio, including territories still missing a boundary."],
        ],
    )
    add_para(
        doc,
        "Financial Commitments by Status groups the same configured amounts by activation status (Activated, Pending Activation, In Progress, Pending Review). Only Activated accounts can receive Confirm Collection.",
    )

    heading(doc, "3.2 Confirm Collection", 2)
    add_para(
        doc,
        "Collections Overview and the collections tables show Activated accounts only. Use Confirm Collection on an Upfront or Billable monthly row to record cash received.",
    )
    steps(
        doc,
        [
            "Stay on Financials Dashboard.",
            "Find an Activated row that still shows Confirm Collection (not the Collected badge).",
            "Click Confirm Collection.",
            "For upfront: enter the amount collected. Partial payments are saved. The row is marked Collected only when the full upfront amount is covered.",
            "For monthly: choose the start period if needed, enter the amount, then confirm. Extra cash can roll into later months.",
            "Click Confirm Collection in the dialog. Cancel leaves the ledger unchanged.",
        ],
    )
    shot(doc, "admin-03-dashboard-confirm.png", "Confirm Collection dialog — amount collected, remaining due, and applied preview.")
    callout(
        doc,
        "One collection ledger",
        "Dashboard, Client Details, Transactions, Revenue, and Reports all read and write esarisari_franchise_collections. Confirm Collection on any of those Admin screens updates the others after navigation. Reset Demo Data clears the same ledger.",
        fill="EEF2FF",
    )

    heading(doc, "3.3 Transactions, Revenue, and Reports", 2)
    add_para(
        doc,
        "Admin can confirm the same franchise setup collections on Transactions, Revenue, and Reports. Those pages keep Internet Credits and sales as a separate stream. Collection cash is not added into Total earnings, and collection rows are not mixed into the sales table.",
    )
    add_table(
        doc,
        ["Page", "What Admin sees"],
        [
            ["Transactions", "Franchise Setup Collections card above the sales table. Shares the page date range. Confirm Collection on unpaid or partial rows."],
            ["Revenue", "A fourth Franchise collections card (cash collected in the period) plus the same collections table. Internet Credits, Sales Commission, and Total earnings stay IC + sales only."],
            ["Reports", "A Franchise collections hero card, Collections by client rollup, the same line-item table, and a Franchise collections CSV export."],
        ],
    )
    steps(
        doc,
        [
            "Confirm a partial upfront or monthly amount on Financials Dashboard.",
            "Open Client Details — paid and remaining match.",
            "Open Transactions, Revenue, or Reports (Admin). Choose All Time or a range that includes the payment date / month.",
            "Confirm the remainder from Revenue (or Transactions / Reports). Enter an amount and a reference number.",
            "Return to Dashboard and Client Details. The same paid / remaining values appear without a second write path.",
        ],
    )
    shot(
        doc,
        "admin-14-transactions-collections.png",
        "Admin Transactions — Franchise Setup Collections sit above the sales table and share the page date range.",
    )
    shot(
        doc,
        "admin-15-revenue-collections.png",
        "Admin Revenue — Franchise collections card plus the same collections table. Total earnings stays Internet Credits + sales.",
    )
    shot(
        doc,
        "admin-16-reports-collections.png",
        "Admin Reports — Franchise collections hero card, Collections by client, and the Franchise collections CSV export.",
    )
    callout(
        doc,
        "What is included",
        "Activated clients only. Upfront = package fees + enabled one-time fees. Billable monthly = fixed monthly fees that are not cost deductions. Company vs client % is informational on the table. % of gross sales, cost-deduction-only fees, and Client Details demo Gross Sale / Cost Deduction / Revenue Share Payout rows stay off these pages.",
        fill="ECFDF5",
    )

    heading(doc, "4. Clients", 1)
    add_para(
        doc,
        "Open Clients in the sidebar. This is the franchise / sub-franchise portfolio: seeded demo rows plus any clients you create with Add New Client. Newly registered clients appear at the top as Pending Review. Seeded examples include Surigao City Service Hub (Activated Sub-Franchisor) and Siargao General Luna Operator (Franchisee, still in progress).",
    )
    shot(doc, "admin-04-clients.png", "Clients list — newly registered clients appear first, then the seeded demo rows.")

    heading(doc, "4.1 Client list", 2)
    add_table(
        doc,
        ["Column", "Meaning"],
        [
            ["Client", "Display name and last-updated time."],
            ["Type", "Sub-Franchisor or Franchisee."],
            ["Status", "Activated, Pending Activation, Pending Review, or In Progress."],
            ["Territories", "Count of assigned coverage territories."],
            ["Upfront Setup", "Package unit fees plus enabled one-time fees."],
            ["Billable Monthly", "Fixed monthly fees that are billed (excludes cost-deduction-only and % gross sales)."],
            ["Revenue Split", "Company % + this client %. A valid split shows 100%. Downline shares are not set by admin."],
        ],
    )
    steps(
        doc,
        [
            "Click View on a row to open Client Details.",
            "Click Add New Client to start the onboarding wizard at Client Info (step 1).",
        ],
    )

    heading(doc, "4.2 Client Details", 2)
    add_para(
        doc,
        "Client Details shows setup, fees, territories, the revenue-split math, and — for Activated clients — a financial history you can collect against. Clients you registered also show a Company Profile card (admin, company, contact, and address from step 1). Seeded demo rows do not have that card.",
    )
    shot(doc, "admin-05-client-detail.png", "Client Details for an Activated account — setup totals, split, and history.")
    bullets(
        doc,
        [
            "Upfront Setup, Billable Fixed Monthly, and Revenue Split cards summarize the configured deal.",
            "Revenue Split Breakdown starts from a demo gross sale, subtracts monthly fees tagged as cost deductions, then applies the company vs this-client split to net revenue.",
            "Financial History (Activated only) lists unpaid and collected items. Confirm Collection here requires an amount and a payment reference.",
            "Use the date filters to narrow history. Back to Clients returns to the list.",
        ],
    )

    heading(doc, "4.3 Activate a client", 2)
    add_para(
        doc,
        "Activation happens on Client Details. There is no separate activation page. Newly registered clients start as Pending Review. Seeded rows that are In Progress or Pending Review can be activated the same way.",
    )
    steps(
        doc,
        [
            "Open Clients and click View on a row that is not Activated.",
            "Click Activate Client in the page header, or in the Financial History empty state.",
            "Confirm Activate Client in the dialog.",
            "Status becomes Activated. Financial History appears, and Confirm Collection is enabled on this page, the Financials Dashboard, and the Admin collections section on Transactions, Revenue, and Reports.",
        ],
    )
    shot(
        doc,
        "admin-13-activate-client.png",
        "Activate Client confirmation — marks the account live without creating a login.",
    )
    callout(
        doc,
        "What activation does",
        "Activate Client saves Activated plus an activated-at timestamp in this browser. It does not create a sign-in account. Reset Demo Data removes activation overlays and registered clients; seeded demo rows return to their original statuses.",
        fill="ECFDF5",
    )

    heading(doc, "5. Add a new client (onboarding)", 1)
    add_para(
        doc,
        "Onboarding is four steps. Each step saves a draft to localStorage as you type or change a control, so Back and Continue keep your place. Create Franchisee on step 4 writes the finished client into the Clients list. Reset Demo Data clears the draft and removes registered clients.",
    )
    add_table(
        doc,
        ["Step", "Screen", "Saves"],
        [
            ["1", "Client Info", "Admin login, company profile, optional contact person"],
            ["2", "Franchise Setup", "Client type, packages, one-time fees, monthly fees, territory"],
            ["3", "Revenue Split", "Company vs this client only. Downline shares are not set here."],
            ["4", "Review", "Read-only summary, then Create Franchisee writes the client to the Clients list"],
        ],
    )
    add_para(
        doc,
        "Start from Clients → Add New Client, or from Financials Dashboard → Open Franchise Setup. Both open step 1.",
    )

    heading(doc, "5.1 Step 1 — Client Info", 2)
    add_para(
        doc,
        "Enter the client admin login, the legal company profile, and an optional operations contact. Continue to Step 2 is blocked until required fields pass validation.",
    )
    shot(doc, "admin-06-onboarding-step-1.png", "Client Info — Admin Profile, Company Profile, and Contact Person.")
    add_table(
        doc,
        ["Section", "Required", "Notes"],
        [
            ["Admin Profile", "First name, last name, email, password, confirm password", "Password must be at least 8 characters and must match confirm."],
            ["Company Profile", "Name, registration number, corporate email, address 1, address 2, city, state/province/region, country, postal/ZIP", "Postal/ZIP is 1–5 digits. Tax ID and company phone are optional."],
            ["Contact Person", "None", "Full name, email, and phone are optional."],
        ],
    )
    steps(
        doc,
        [
            "Fill Admin Profile and Company Profile.",
            "Add a Contact Person if you have one.",
            "Click Continue to Step 2. Fix any red field messages before the wizard advances.",
        ],
    )

    heading(doc, "5.2 Step 2 — Franchise Setup", 2)
    add_para(
        doc,
        "Choose whether this client is a Sub-Franchisor or a Franchisee, then configure packages, fees, and coverage.",
    )
    shot(doc, "admin-07-onboarding-step-2.png", "Franchise Setup — client type, packages, fees, and territory.")
    bullets(
        doc,
        [
            "Client type: Sub-Franchisor oversees franchisees in a territory. Franchisee operates units. This choice only changes the client label on step 3 (Sub-franchisor vs Franchisee).",
            "Packages: eNeighborhood (₱20,000 / unit), eBarangay (₱200,000 / unit), eLGU (₱2,000,000 / unit). Set quantity and a primary package.",
            "One-time fees default to Franchise Fee, Setup & Training, and Legal & Documentation. Toggle or edit amounts. Enabled fees roll into upfront due.",
            "Monthly operational fees can be Fixed monthly, % of gross sales, or cost deductions. Billable monthly on Dashboard / Clients excludes cost-deduction-only and % gross sales lines.",
            "Territory and areas assign coverage. A missing boundary is flagged on the Dashboard Coverage card for seeded clients.",
        ],
    )
    steps(
        doc,
        [
            "Select Sub-Franchisor or Franchisee.",
            "Adjust package quantities and the primary package if needed.",
            "Review one-time and monthly fees.",
            "Select a territory and areas.",
            "Click Continue to Step 3. Use Back to Client Info if you need to change the company record.",
        ],
    )

    heading(doc, "5.3 Step 3 — Revenue Split", 2)
    add_para(
        doc,
        "Admin sets only the contract between eSariSari (company) and this client. Changing one percentage fills the other so the two always add up to 100%. Downline shares are not editable here.",
    )
    shot(doc, "admin-08-onboarding-step-3.png", "Revenue Split — company vs this client only.")
    add_table(
        doc,
        ["Client type", "Admin sets", "Not set by admin"],
        [
            ["Sub-Franchisor", "Company % and this sub-franchisor %", "Franchisee and retailer shares (sub-franchisor → franchisee → retailers)"],
            ["Franchisee", "Company % and this franchisee %", "Retailer shares (franchisee → retailers)"],
        ],
    )
    steps(
        doc,
        [
            "Edit Company or the client share. The other value updates so the total stays 100%.",
            "Default is Company 40% / client 60%. Use Reset defaults to restore that.",
            "Click Continue to Review.",
        ],
    )

    heading(doc, "5.4 Step 4 — Review", 2)
    add_para(
        doc,
        "Review is read-only. Each card has an Edit link back to the step that owns those fields. Check client type, admin and company profile, packages, territory, fees, and the split bar.",
    )
    shot(doc, "admin-09-onboarding-step-4.png", "Review — confirm the draft, then Create Franchisee.")
    steps(
        doc,
        [
            "Scan every card. Use Edit if a value is wrong.",
            "Click Create Franchisee.",
            "If Client Info is incomplete or the split is not 100%, fix the message and try again.",
            "The app opens Clients. Your new row is at the top as Pending Review.",
            "Click View on that row to confirm Company Profile and the rest of the setup.",
        ],
    )
    shot(
        doc,
        "admin-11-registered-client.png",
        "Clients list after Create Franchisee — the new company is first, Pending Review.",
    )
    shot(
        doc,
        "admin-12-registered-client-detail.png",
        "Client Details for a registered client — Company Profile from step 1.",
    )
    callout(
        doc,
        "What gets saved",
        "Create Franchisee writes the company record to this browser (esarisari_registered_clients) and opens Clients. Passwords are not stored. This does not create a sign-in account for the new client. Open View, then Activate Client, before Confirm Collection. Reset Demo Data removes registered clients, activation overlays, and the onboarding draft; seeded demo clients stay.",
        fill="ECFDF5",
    )

    heading(doc, "6. Reset Demo Data and logout", 1)
    shot(doc, "admin-10-user-menu.png", "Name menu — Profile, both user guides, Reset Demo Data, and Logout.")
    add_para(
        doc,
        "Open your name at the top right. Admin sees User Guide (Internet Credits) and Admin User Guide (this document), plus Download computations (Excel), Reset Demo Data, Profile, and Logout.",
    )
    heading(doc, "6.1 Reset Demo Data", 2)
    add_para(
        doc,
        "Reset restores the seeded network, wallets, empty Internet Credits / sales history, the seeded client portfolio, and collection marks. It also removes clients you registered, clears activation overlays, and clears the onboarding draft (client info, franchise setup, and revenue split).",
    )
    steps(
        doc,
        [
            "Click your name at the top right.",
            "Click Reset Demo Data.",
            "Confirm Reset Demo Data in the dialog.",
        ],
    )
    callout(
        doc,
        "Training tip",
        "Reset before a live walkthrough so Dashboard collections, the Clients list, and the onboarding draft match this guide. Do not reset mid-demo unless you intend to wipe the session — including any clients you just registered.",
        fill="EEF2FF",
    )

    heading(doc, "6.2 Logout", 2)
    steps(
        doc,
        [
            "Click your name at the top right.",
            "Click Logout.",
            "Sign in again as Admin, or use another demo account for the Internet Credits guide.",
        ],
    )

    heading(doc, "7. Field meanings", 1)
    add_table(
        doc,
        ["Term", "Meaning"],
        [
            ["Activated", "Client is live. Dashboard, Client Details, and the Admin collections section on Transactions, Revenue, and Reports can Confirm Collection. Use Activate Client on Client Details to get here."],
            ["Upfront / one-time", "Package unit fees plus enabled one-time franchise fees."],
            ["Billable fixed monthly", "Monthly fees billed as a peso amount. Excludes cost-deduction-only and % of gross sales."],
            ["Cost deduction", "A monthly fee subtracted from gross sales before the revenue split."],
            ["% of gross sales", "A fee expressed as a percent of sales. Tracked on the client record; not in the billable monthly KPI."],
            ["Revenue split", "Company + this client. Must equal 100%. Downline retailer / franchisee shares are set by the client."],
            ["Net revenue for sharing", "Demo gross sale minus standard cost deductions."],
        ],
    )

    heading(doc, "8. Common mistakes", 1)
    bullets(
        doc,
        [
            "Using password123 on the Admin account — Admin is Restricted and uses abc12345678.",
            "Looking for Franchises or Organizations in the sidebar — those items were removed. Use Clients.",
            "Expecting franchise collections inside Total earnings or the sales table — those stay Internet Credits and sales. Admin collections are a separate section on Transactions, Revenue, and Reports.",
            "Expecting Confirm Collection on a non-Activated client — open Client Details and click Activate Client first.",
            "Treating Create Franchisee as a login for the new client — the company appears on Clients, but this demo does not create a sign-in account.",
            "Expecting the new client to stay off the Clients list — Create Franchisee now saves it and shows it at the top as Pending Review.",
            "Trying to set retailer or franchisee downline shares on step 3 — those belong to the client, not platform admin.",
            "Mixing this guide with Internet Credits work — loads, sales, and commissions are in the other User Guide.",
            "Using Reset Demo Data during a live collection walkthrough — it wipes payment marks, the onboarding draft, registered clients, and activation overlays.",
        ],
    )

    add_para(
        doc,
        f"Admin User Guide version {GUIDE_VERSION}  ·  {GUIDE_DATE}. Matches the eSariSari Admin module on that date. It does not replace the Internet Credits User Guide (version 2.0). Screenshots are layout references; live numbers depend on collections you confirm in this session.",
        italic=True,
        color=MUTED,
        space_before=12,
    )

    doc.save(OUTPUT)
    OUTPUT_HTML.write_text(HTML.render(), encoding="utf-8")
    print(f"Wrote {OUTPUT}")
    print(f"Wrote {OUTPUT_HTML}")


if __name__ == "__main__":
    build()
