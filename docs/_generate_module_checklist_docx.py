#!/usr/bin/env python3
"""Generate docs/module-checklist-by-role.docx."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "module-checklist-by-role.docx"

NAVY = RGBColor(0x0F, 0x17, 0x2A)
BLUE = RGBColor(0x1D, 0x4E, 0xD8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
MUTED = RGBColor(0x47, 0x55, 0x69)
DONE = RGBColor(0x16, 0x65, 0x34)
WIP = RGBColor(0xB4, 0x53, 0x09)
TODO = RGBColor(0x9F, 0x12, 0x39)

DONE_BG = "DCFCE7"
WIP_BG = "FEF3C7"
TODO_BG = "FEE2E2"

STATUS_COLOR = {
    "Completed": DONE,
    "In progress": WIP,
    "To do": TODO,
}
STATUS_BG = {
    "Completed": DONE_BG,
    "In progress": WIP_BG,
    "To do": TODO_BG,
}


def set_run(run, *, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    r = run._element.get_or_add_rPr()
    rFonts = r.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        r.append(rFonts)
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rFonts.set(qn("w:eastAsia"), "Calibri")


def shade_cell(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_borders(cell, color="E2E8F0"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)


def v_align(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    val = tcPr.find(qn("w:vAlign"))
    if val is None:
        val = OxmlElement("w:vAlign")
        tcPr.append(val)
    val.set(qn("w:val"), "center")


def set_cell_text(cell, text, *, size=10, bold=False, color=NAVY):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)


def add_para(
    doc, text, *, size=11, bold=False, color=None, space_after=8, italic=False
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color or NAVY, italic=italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    return p


def task_row(title, subtitle, status):
    return (title, subtitle, status)


def live_page_tasks(
    *,
    ux_note,
    fe_note,
    api_ready="In progress",
    api_ready_note="No client yet for this role",
    wiring="In progress",
    wiring_note="Still on localStorage",
    qa="To do",
    qa_note="Pending verification",
    docs="Completed",
    docs_note="Covered in app-modules / user guide",
):
    return [
        task_row("Initial UI/UX", ux_note, "Completed"),
        task_row("Frontend implementation", fe_note, "Completed"),
        task_row("API-ready client", api_ready_note, api_ready),
        task_row("API wiring", wiring_note, wiring),
        task_row("QA / verification", qa_note, qa),
        task_row("Documentation", docs_note, docs),
    ]


def placeholder_page_tasks(*, ux_note, fe_note):
    return [
        task_row("Initial UI/UX", ux_note, "To do"),
        task_row("Frontend implementation", fe_note, "To do"),
        task_row("API-ready client", "No contract yet", "In progress"),
        task_row("API wiring", "Blocked on frontend + API", "In progress"),
        task_row("QA / verification", "Page is not live", "To do"),
        task_row("Documentation", "Placeholder only in app-modules", "In progress"),
    ]


SUB_API = dict(
    api_ready="In progress",
    api_ready_note="Unwired client in src/services/api/subfranchisor",
    wiring="In progress",
    wiring_note="Do not swap loaders until VITE_API_BASE_URL is set",
)

ADMIN_API = dict(
    api_ready="In progress",
    api_ready_note="Need /api/v1 admin + finops resources",
    wiring="In progress",
    wiring_note="Still on localStorage",
)

FRAN_API = dict(
    api_ready="In progress",
    api_ready_note="Need franchisee API module (same pattern as Sub)",
    wiring="In progress",
    wiring_note="Still on localStorage",
)

RETAIL_API = dict(
    api_ready="In progress",
    api_ready_note="Need retailer API module",
    wiring="In progress",
    wiring_note="Still on localStorage",
)


ROLES = [
    {
        "title": "Shared shell",
        "subtitle": "All authenticated roles",
        "intro": "Chrome that every role uses. Not a sidebar module, but it must be tracked with the rest.",
        "modules": [
            {
                "name": "Authentication & chrome",
                "pages": [
                    {
                        "name": "Login",
                        "path": "/login",
                        "tasks": live_page_tasks(
                            ux_note="Email, password, demo account chips",
                            fe_note="Session in localStorage; admin password restricted",
                            **ADMIN_API,
                            docs_note="Documented in app-modules §3",
                        ),
                    },
                    {
                        "name": "Profile",
                        "path": "/profile",
                        "tasks": live_page_tasks(
                            ux_note="Read-only name, email, role, organization",
                            fe_note="Reads AuthContext only; no edit",
                            **ADMIN_API,
                        ),
                    },
                    {
                        "name": "Notifications",
                        "path": "Header bell",
                        "tasks": live_page_tasks(
                            ux_note="Pending credit requests + low/zero credits",
                            fe_note="Per-org read state in localStorage",
                            api_ready="In progress",
                            api_ready_note="Sub client has notifications.js; other roles do not",
                            wiring="In progress",
                            wiring_note="Bell still uses storage.js",
                        ),
                    },
                    {
                        "name": "Reset Demo Data",
                        "path": "User menu",
                        "tasks": [
                            task_row(
                                "Initial UI/UX",
                                "Confirm in the user menu",
                                "Completed",
                            ),
                            task_row(
                                "Frontend implementation",
                                "Reseeds localStorage; keeps session if user still exists",
                                "Completed",
                            ),
                            task_row(
                                "API-ready client",
                                "Demo-only; not a production API",
                                "In progress",
                            ),
                            task_row(
                                "API wiring",
                                "Not applicable until a backend seed/reset exists",
                                "In progress",
                            ),
                            task_row(
                                "QA / verification",
                                "Pending verification",
                                "To do",
                            ),
                            task_row(
                                "Documentation",
                                "Called out in app-modules and user guides",
                                "Completed",
                            ),
                        ],
                    },
                ],
            }
        ],
    },
    {
        "title": "Admin module",
        "subtitle": "Platform Admin — home: /dashboard",
        "intro": "Franchise setup collections plus the network Internet Credits / sales ledger. Live UI on localStorage. API-ready client and API wiring are in progress.",
        "modules": [
            {
                "name": "Financials Dashboard",
                "pages": [
                    {
                        "name": "Dashboard",
                        "path": "/dashboard",
                        "tasks": live_page_tasks(
                            ux_note="Upfront + billable monthly collections KPIs and tables",
                            fe_note="esarisari_franchise_collections ledger",
                            **ADMIN_API,
                        ),
                    }
                ],
            },
            {
                "name": "Clients & franchise setup",
                "pages": [
                    {
                        "name": "Clients list",
                        "path": "/franchise-setup/clients",
                        "tasks": live_page_tasks(
                            ux_note="Registered + seeded portfolio; Add New Client",
                            fe_note="localStorage registered clients + demo portfolio",
                            **ADMIN_API,
                        ),
                    },
                    {
                        "name": "Client details",
                        "path": "/franchise-setup/clients/:clientId",
                        "tasks": live_page_tasks(
                            ux_note="Activate, split, collections, territory",
                            fe_note="Same collection ledger as Dashboard",
                            **ADMIN_API,
                        ),
                    },
                    {
                        "name": "Onboarding step 1 — Client Info",
                        "path": "/franchise-setup/onboarding/step-1",
                        "tasks": live_page_tasks(
                            ux_note="Admin credentials + company profile",
                            fe_note="Draft in esarisari_onboarding_client_info",
                            **ADMIN_API,
                        ),
                    },
                    {
                        "name": "Onboarding step 2 — Franchise Setup",
                        "path": "/franchise-setup/onboarding/step-2",
                        "tasks": live_page_tasks(
                            ux_note="Client type, packages, territory, fees",
                            fe_note="Draft in esarisari_onboarding_franchise_setup",
                            **ADMIN_API,
                        ),
                    },
                    {
                        "name": "Onboarding step 3 — Revenue Split",
                        "path": "/franchise-setup/onboarding/step-3",
                        "tasks": live_page_tasks(
                            ux_note="Company vs this client only; downlines out of scope",
                            fe_note="esarisari_onboarding_revenue_split",
                            **ADMIN_API,
                        ),
                    },
                    {
                        "name": "Onboarding step 4 — Review",
                        "path": "/franchise-setup/onboarding/step-4",
                        "tasks": live_page_tasks(
                            ux_note="Confirm and Create Franchisee",
                            fe_note="Writes esarisari_registered_clients",
                            **ADMIN_API,
                        ),
                    },
                ],
            },
            {
                "name": "Wallets",
                "pages": [
                    {
                        "name": "Wallets directory",
                        "path": "/wallets",
                        "tasks": live_page_tasks(
                            ux_note="Network Available Credits directory",
                            fe_note="Operating wallets from localStorage",
                            **ADMIN_API,
                        ),
                    }
                ],
            },
            {
                "name": "Internet Credits",
                "pages": [
                    {
                        "name": "Internet Credits workspace",
                        "path": "/funding",
                        "tasks": live_page_tasks(
                            ux_note="Incoming, mine, direct release, approve/reject",
                            fe_note="fundingActions.js + localStorage ledgers",
                            **ADMIN_API,
                        ),
                    }
                ],
            },
            {
                "name": "Deposit Rates",
                "pages": [
                    {
                        "name": "Deposit Rates",
                        "path": "/deposit-rates",
                        "tasks": live_page_tasks(
                            ux_note="Admin → Sub hop overrides",
                            fe_note="esarisari_deposit_rates",
                            **ADMIN_API,
                        ),
                    }
                ],
            },
            {
                "name": "Commission Settings",
                "pages": [
                    {
                        "name": "Commission Settings",
                        "path": "/commission-settings",
                        "tasks": live_page_tasks(
                            ux_note="Per-retailer split; Admin-editable platform fee",
                            fe_note="% of sales; Sub share is remainder for Admin",
                            **ADMIN_API,
                        ),
                    }
                ],
            },
            {
                "name": "Transactions",
                "pages": [
                    {
                        "name": "Transactions Ledger",
                        "path": "/transactions",
                        "tasks": live_page_tasks(
                            ux_note="Sales table + franchise setup collections panel",
                            fe_note="Empty sales seed; collections share Dashboard ledger",
                            **ADMIN_API,
                        ),
                    }
                ],
            },
            {
                "name": "Revenue",
                "pages": [
                    {
                        "name": "Revenue",
                        "path": "/revenue",
                        "tasks": live_page_tasks(
                            ux_note="IC earnings + sales commission + collections",
                            fe_note="Two money streams stay separate from collections cash",
                            **ADMIN_API,
                        ),
                    }
                ],
            },
            {
                "name": "Reports",
                "pages": [
                    {
                        "name": "Reports",
                        "path": "/reports",
                        "tasks": live_page_tasks(
                            ux_note="Network tables, Sub revenue-sharing sheet, CSV, collections",
                            fe_note="Client-sheet Sub report included for Admin too",
                            **ADMIN_API,
                        ),
                    }
                ],
            },
            {
                "name": "Leftovers (not in sidebar)",
                "pages": [
                    {
                        "name": "Organizations",
                        "path": "/organizations",
                        "tasks": placeholder_page_tasks(
                            ux_note="Placeholder copy; removed from nav",
                            fe_note="Route exists; not a product surface",
                        ),
                    }
                ],
            },
        ],
    },
    {
        "title": "Sub-Franchisee module",
        "subtitle": "Sub-Franchisee — home: /wallet-management",
        "intro": "Operational network: wallets, Internet Credits, rates, commission, sales ledger, reports. Frontend is live on localStorage. API-ready client and API wiring are in progress.",
        "modules": [
            {
                "name": "Dashboard (greyed)",
                "pages": [
                    {
                        "name": "Dashboard",
                        "path": "/dashboard",
                        "tasks": [
                            task_row(
                                "Initial UI/UX",
                                "Greyed nav item only; Admin Financials is not this role",
                                "Completed",
                            ),
                            task_row(
                                "Frontend implementation",
                                "Role-specific Sub dashboard",
                                "To do",
                            ),
                            task_row(
                                "API-ready client",
                                "Not in the current Sub API contract",
                                "In progress",
                            ),
                            task_row(
                                "API wiring",
                                "Blocked on product decision",
                                "In progress",
                            ),
                            task_row(
                                "QA / verification",
                                "Route is blocked for this role",
                                "To do",
                            ),
                            task_row(
                                "Documentation",
                                "Noted as greyed in app-modules",
                                "Completed",
                            ),
                        ],
                    }
                ],
            },
            {
                "name": "Franchisees (greyed)",
                "pages": [
                    {
                        "name": "Franchisees",
                        "path": "/franchisees",
                        "tasks": placeholder_page_tasks(
                            ux_note="Greyed nav; no AppRoutes entry",
                            fe_note="Placeholder page file only",
                        ),
                    }
                ],
            },
            {
                "name": "Wallets",
                "pages": [
                    {
                        "name": "Wallet Management",
                        "path": "/wallet-management",
                        "tasks": live_page_tasks(
                            ux_note="Own org + franchisee + retailer credits directory",
                            fe_note="Read-only; activity sheet from transfers",
                            **SUB_API,
                        ),
                    }
                ],
            },
            {
                "name": "Internet Credits",
                "pages": [
                    {
                        "name": "Internet Credits workspace",
                        "path": "/funding",
                        "tasks": live_page_tasks(
                            ux_note="Request from Admin; release/reject franchisee requests; direct release",
                            fe_note="releaseSource=balance; hop admin_to_sub / sub_to_franchisee",
                            **SUB_API,
                        ),
                    }
                ],
            },
            {
                "name": "Deposit Rates",
                "pages": [
                    {
                        "name": "Deposit Rates",
                        "path": "/deposit-rates",
                        "tasks": live_page_tasks(
                            ux_note="Sub → Franchisee hop (default 70%)",
                            fe_note="Overrides in esarisari_deposit_rates",
                            **SUB_API,
                        ),
                    }
                ],
            },
            {
                "name": "Commission Settings",
                "pages": [
                    {
                        "name": "Commission Settings",
                        "path": "/commission-settings",
                        "tasks": live_page_tasks(
                            ux_note="Retailer, franchisee, Your Share editable; platform fee read-only",
                            fe_note="Split is % of sales; total must equal 100%",
                            **SUB_API,
                        ),
                    }
                ],
            },
            {
                "name": "Transactions",
                "pages": [
                    {
                        "name": "Transactions Ledger",
                        "path": "/transactions",
                        "tasks": live_page_tasks(
                            ux_note="Network sales; Your Share column; no Record sale",
                            fe_note="Scoped by subfranchiseeOrganizationId",
                            **SUB_API,
                        ),
                    }
                ],
            },
            {
                "name": "Revenue",
                "pages": [
                    {
                        "name": "Revenue",
                        "path": "/revenue",
                        "tasks": live_page_tasks(
                            ux_note="IC spread earnings + sales commission",
                            fe_note="credit_spread mode; no franchise collections",
                            **SUB_API,
                        ),
                    }
                ],
            },
            {
                "name": "Reports",
                "pages": [
                    {
                        "name": "Reports",
                        "path": "/reports",
                        "tasks": live_page_tasks(
                            ux_note="Revenue Sharing Sub-Franchisee table + CSV + network commissions",
                            fe_note="Client-sheet layout; platform fee excluded from Total Revenue",
                            **SUB_API,
                        ),
                    }
                ],
            },
        ],
    },
    {
        "title": "Franchisee module",
        "subtitle": "Franchisee — home: /wallet-management",
        "intro": "Same operational surfaces as Sub, scoped to this franchisee’s retailers. API-ready client and API wiring are in progress. Commission Settings is Admin/Sub only.",
        "modules": [
            {
                "name": "Dashboard (greyed)",
                "pages": [
                    {
                        "name": "Dashboard",
                        "path": "/dashboard",
                        "tasks": [
                            task_row(
                                "Initial UI/UX",
                                "Greyed nav item only",
                                "Completed",
                            ),
                            task_row(
                                "Frontend implementation",
                                "Role-specific franchisee dashboard",
                                "To do",
                            ),
                            task_row(
                                "API-ready client",
                                "Not started",
                                "In progress",
                            ),
                            task_row(
                                "API wiring",
                                "Blocked on product + API",
                                "In progress",
                            ),
                            task_row(
                                "QA / verification",
                                "Route is blocked",
                                "To do",
                            ),
                            task_row(
                                "Documentation",
                                "Noted as greyed in app-modules",
                                "Completed",
                            ),
                        ],
                    }
                ],
            },
            {
                "name": "Retailers (greyed)",
                "pages": [
                    {
                        "name": "Retailers",
                        "path": "/retailers",
                        "tasks": placeholder_page_tasks(
                            ux_note="Greyed nav; placeholder page; route blocked",
                            fe_note="No retailer directory product yet",
                        ),
                    }
                ],
            },
            {
                "name": "Wallets",
                "pages": [
                    {
                        "name": "Wallet Management",
                        "path": "/wallet-management",
                        "tasks": live_page_tasks(
                            ux_note="Own + retailer credits directory",
                            fe_note="Read-only localStorage wallets",
                            **FRAN_API,
                        ),
                    }
                ],
            },
            {
                "name": "Internet Credits",
                "pages": [
                    {
                        "name": "Internet Credits workspace",
                        "path": "/funding",
                        "tasks": live_page_tasks(
                            ux_note="Request from Sub; release/reject retailer requests; direct release",
                            fe_note="Hop sub_to_franchisee / franchisee_to_retailer",
                            **FRAN_API,
                        ),
                    }
                ],
            },
            {
                "name": "Deposit Rates",
                "pages": [
                    {
                        "name": "Deposit Rates",
                        "path": "/deposit-rates",
                        "tasks": live_page_tasks(
                            ux_note="Franchisee → Retailer hop (default 80%)",
                            fe_note="Downline overrides only",
                            **FRAN_API,
                        ),
                    }
                ],
            },
            {
                "name": "Transactions",
                "pages": [
                    {
                        "name": "Transactions Ledger",
                        "path": "/transactions",
                        "tasks": live_page_tasks(
                            ux_note="Retailer sales under this franchisee",
                            fe_note="Scoped by franchiseeOrganizationId",
                            **FRAN_API,
                        ),
                    }
                ],
            },
            {
                "name": "Revenue",
                "pages": [
                    {
                        "name": "Revenue",
                        "path": "/revenue",
                        "tasks": live_page_tasks(
                            ux_note="IC spread + sales commission",
                            fe_note="credit_spread mode",
                            **FRAN_API,
                        ),
                    }
                ],
            },
            {
                "name": "Reports",
                "pages": [
                    {
                        "name": "Reports",
                        "path": "/reports",
                        "tasks": live_page_tasks(
                            ux_note="Retailer commissions + IC earnings; no Sub sheet table",
                            fe_note="Role config hides Sub revenue-sharing export",
                            **FRAN_API,
                        ),
                    }
                ],
            },
        ],
    },
    {
        "title": "Retailer module",
        "subtitle": "Retailer — home: /wallet",
        "intro": "Wallet, request credits, record demo sales, and view own commission. API-ready client and API wiring are in progress.",
        "modules": [
            {
                "name": "Dashboard (greyed)",
                "pages": [
                    {
                        "name": "Dashboard",
                        "path": "/dashboard",
                        "tasks": [
                            task_row(
                                "Initial UI/UX",
                                "Greyed nav item only",
                                "Completed",
                            ),
                            task_row(
                                "Frontend implementation",
                                "Role-specific retailer dashboard",
                                "To do",
                            ),
                            task_row(
                                "API-ready client",
                                "Not started",
                                "In progress",
                            ),
                            task_row(
                                "API wiring",
                                "Blocked on product + API",
                                "In progress",
                            ),
                            task_row(
                                "QA / verification",
                                "Route is blocked",
                                "To do",
                            ),
                            task_row(
                                "Documentation",
                                "Noted as greyed in app-modules",
                                "Completed",
                            ),
                        ],
                    }
                ],
            },
            {
                "name": "Wallet",
                "pages": [
                    {
                        "name": "Wallet",
                        "path": "/wallet",
                        "tasks": live_page_tasks(
                            ux_note="Available Credits, sale margin, request CTA",
                            fe_note="Retailer operating wallet; cannot record sales here",
                            **RETAIL_API,
                        ),
                    }
                ],
            },
            {
                "name": "Internet Credits",
                "pages": [
                    {
                        "name": "Request Funding",
                        "path": "/request-funding",
                        "tasks": live_page_tasks(
                            ux_note="Request-only workspace to Franchisee",
                            fe_note="Same FundingWorkspacePage; no incoming / direct release",
                            **RETAIL_API,
                        ),
                    }
                ],
            },
            {
                "name": "Transactions",
                "pages": [
                    {
                        "name": "Transactions Ledger",
                        "path": "/transactions",
                        "tasks": live_page_tasks(
                            ux_note="Own sales + Record demo sale",
                            fe_note="Burns ~97% credits; stamps commission % of sales",
                            **RETAIL_API,
                        ),
                    }
                ],
            },
            {
                "name": "Revenue",
                "pages": [
                    {
                        "name": "Revenue",
                        "path": "/revenue",
                        "tasks": live_page_tasks(
                            ux_note="Your commission, sales volume, credits consumed",
                            fe_note="No IC earnings / Total earnings cards",
                            **RETAIL_API,
                        ),
                    }
                ],
            },
            {
                "name": "Reports",
                "pages": [
                    {
                        "name": "Reports",
                        "path": "/reports",
                        "tasks": live_page_tasks(
                            ux_note="Own sales commission + volume; limited CSV",
                            fe_note="No network tables or Sub sheet",
                            **RETAIL_API,
                        ),
                    }
                ],
            },
        ],
    },
]


def count_statuses():
    counts = {"Completed": 0, "In progress": 0, "To do": 0}
    for role in ROLES:
        for module in role["modules"]:
            for page in module["pages"]:
                for _title, _sub, status in page["tasks"]:
                    counts[status] = counts.get(status, 0) + 1
    return counts


def add_task_table(doc, tasks):
    table = doc.add_table(rows=1 + len(tasks), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    headers = ["Task title", "Subtitle", "Status"]
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "0F172A")
        set_cell_borders(cell, "0F172A")
        v_align(cell)
        set_cell_text(cell, label, size=9, bold=True, color=WHITE)
    for ri, (title, subtitle, status) in enumerate(tasks):
        row = table.rows[ri + 1]
        values = [f"{title} — {subtitle}" if False else title, subtitle, status]
        for ci, value in enumerate(values):
            cell = row.cells[ci]
            fill = "F8FAFC" if ri % 2 == 0 else "FFFFFF"
            if ci == 2:
                fill = STATUS_BG.get(status, fill)
            shade_cell(cell, fill)
            set_cell_borders(cell, "E2E8F0")
            v_align(cell)
            color = STATUS_COLOR.get(status, NAVY) if ci == 2 else NAVY
            set_cell_text(
                cell,
                value,
                size=9,
                bold=ci == 2,
                color=color,
            )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def add_line_items(doc, tasks):
    for title, subtitle, status in tasks:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.15)
        run = p.add_run(f"{title} — {subtitle}  —  {status}")
        set_run(run, size=10.5, color=NAVY)
        # Re-color just the status by splitting is messy; keep table as source of truth.
    # User asked for line format; table is clearer. We'll print lines AND keep table? 
    # They explicitly wanted: <task title - subtitle> - <status>
    # I'll use line items as primary and skip the table to match the request.


def add_checklist_lines(doc, tasks):
    for title, subtitle, status in tasks:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.2)
        run = p.add_run(f"{title} — {subtitle}")
        set_run(run, size=11, color=NAVY)
        run2 = p.add_run("  —  ")
        set_run(run2, size=11, color=MUTED)
        run3 = p.add_run(status)
        set_run(run3, size=11, bold=True, color=STATUS_COLOR.get(status, NAVY))


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("eSariSari FinOps — Module checklist by role")
    set_run(run, size=22, bold=True, color=NAVY)
    title.paragraph_format.space_after = Pt(4)

    add_para(
        doc,
        "1 September 2026  ·  Status against the current localStorage demo. "
        "API-ready client and API wiring are in progress. QA / verification is to do.",
        size=11,
        color=MUTED,
        space_after=10,
    )
    add_para(
        doc,
        "Format: each role is a module group. Under a page, each line is  "
        "Task title — subtitle  —  Status (Completed, In progress, To do).",
        size=11,
        space_after=10,
    )

    counts = count_statuses()
    add_para(
        doc,
        f"Roll-up: {counts['Completed']} completed  ·  "
        f"{counts['In progress']} in progress  ·  {counts['To do']} to do.",
        size=11,
        bold=True,
        space_after=6,
    )

    heading(doc, "How to read status", 1)
    add_para(
        doc,
        "Completed — live in the demo (UI and/or localStorage frontend, or documented).",
        size=11,
        space_after=4,
    )
    add_para(
        doc,
        "In progress — started but not finished (API-ready client and API wiring).",
        size=11,
        space_after=4,
    )
    add_para(
        doc,
        "To do — not started, greyed/placeholder, or blocked on a later wiring task.",
        size=11,
        space_after=12,
    )
    add_para(
        doc,
        "Standard tasks per page: Initial UI/UX · Frontend implementation · "
        "API-ready client · API wiring · QA / verification · Documentation.",
        size=11,
        space_after=12,
    )

    for role in ROLES:
        heading(doc, role["title"], 1)
        add_para(doc, role["subtitle"], size=11, italic=True, color=MUTED, space_after=6)
        add_para(doc, role["intro"], size=11, space_after=10)

        for module in role["modules"]:
            heading(doc, module["name"], 2)
            for page in module["pages"]:
                heading(doc, f"{page['name']}  ·  {page['path']}", 3)
                add_checklist_lines(doc, page["tasks"])

    heading(doc, "Suggested next wiring order", 1)
    add_para(
        doc,
        "1. Sub-Franchisee module — client already exists; set VITE_API_BASE_URL and swap one resource at a time (auth → wallets → internet credits → rates → commission → ledger → reports).",
        size=11,
        space_after=4,
    )
    add_para(
        doc,
        "2. Franchisee module — clone the Sub API pattern, scoped to franchisee downlines.",
        size=11,
        space_after=4,
    )
    add_para(
        doc,
        "3. Retailer module — wallet, request credits, record sale, own revenue/reports.",
        size=11,
        space_after=4,
    )
    add_para(
        doc,
        "4. Admin module — franchise setup collections plus network IC/sales; largest surface.",
        size=11,
        space_after=4,
    )
    add_para(
        doc,
        "Leave greyed Dashboard / Franchisees / Retailers until product wants those pages live.",
        size=11,
        space_after=12,
    )

    add_para(
        doc,
        "Related: docs/app-modules.md  ·  docs/api-subfranchisee.md",
        size=10,
        color=MUTED,
        space_after=0,
    )

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    print(
        "Counts:",
        count_statuses(),
    )


if __name__ == "__main__":
    build()
