#!/usr/bin/env python3
"""Generate a short client-facing module checklist (docx)."""

from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "module-checklist-client.docx"

SKIP_MODULES = {"Leftovers (not in sidebar)"}
SKIP_PAGES = {"Reset Demo Data"}

TASK_COLUMNS = [
    ("Initial UI/UX", "UI / UX"),
    ("Frontend implementation", "Frontend"),
    ("API-ready client", "API client"),
    ("API wiring", "API wiring"),
    ("QA / verification", "QA/UAT"),
    ("Documentation", "Docs"),
]

NAVY = RGBColor(0x0F, 0x17, 0x2A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
MUTED = RGBColor(0x47, 0x55, 0x69)
DONE = RGBColor(0x16, 0x65, 0x34)
WIP = RGBColor(0xB4, 0x53, 0x09)
TODO = RGBColor(0x9F, 0x12, 0x39)

STATUS_BG = {
    "Completed": "DCFCE7",
    "In progress": "FEF3C7",
    "To do": "FEE2E2",
}
STATUS_COLOR = {
    "Completed": DONE,
    "In progress": WIP,
    "To do": TODO,
}
STATUS_SHORT = {
    "Completed": "Completed",
    "In progress": "In progress",
    "To do": "To do",
}

ROLE_LABEL = {
    "Shared shell": "Shared (all roles)",
    "Admin module": "Admin",
    "Sub-Franchisee module": "Sub-Franchisee",
    "Franchisee module": "Franchisee",
    "Retailer module": "Retailer",
}

ROLE_BLURB = {
    "Shared shell": "Sign-in, profile, and notifications used by every role.",
    "Admin module": "Franchise setup, collections, network credits, and sales ledger.",
    "Sub-Franchisee module": "Wallets, credits, rates, commission, sales, and reports.",
    "Franchisee module": "Same operational pages as Sub, scoped to this franchisee’s retailers.",
    "Retailer module": "Wallet, request credits, record sales, and own commission.",
}


def load_checklist():
    path = ROOT / "_generate_module_checklist_docx.py"
    spec = importlib.util.spec_from_file_location("module_checklist", path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Cannot load {path}")
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def set_run(run, *, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    r = run._element.get_or_add_rPr()
    r_fonts = r.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Calibri")


def shade_cell(cell, fill):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_borders(cell, color="E2E8F0"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)


def v_align(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    val = tc_pr.find(qn("w:vAlign"))
    if val is None:
        val = OxmlElement("w:vAlign")
        tc_pr.append(val)
    val.set(qn("w:val"), "center")


def set_cell_width(cell, inches):
    cell.width = Inches(inches)
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    twips = str(int(inches * 1440))
    tc_w.set(qn("w:w"), twips)
    tc_w.set(qn("w:type"), "dxa")


def set_cell_text(
    cell,
    text,
    *,
    size=9,
    bold=False,
    color=NAVY,
    align="left",
):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
    }[align]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)


def add_para(doc, text, *, size=11, bold=False, color=None, space_after=8, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color or NAVY, italic=italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    return p


def clean_module_name(name: str) -> str:
    return name.replace(" (greyed)", " (not live)")


def page_rows(roles):
    rows = []
    counts = {"Completed": 0, "In progress": 0, "To do": 0}
    for role in roles:
        for module in role["modules"]:
            if module["name"] in SKIP_MODULES:
                continue
            for page in module["pages"]:
                if page["name"] in SKIP_PAGES:
                    continue
                statuses = {title: status for title, _sub, status in page["tasks"]}
                for _title, _sub, status in page["tasks"]:
                    counts[status] = counts.get(status, 0) + 1
                rows.append(
                    {
                        "role": role["title"],
                        "module": clean_module_name(module["name"]),
                        "page": page["name"],
                        "statuses": statuses,
                    }
                )
    return rows, counts


def set_table_full_width(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")


def add_role_table(doc, pages):
    headers = ["Module", "Page"] + [label for _key, label in TASK_COLUMNS]
    widths = [1.55, 1.7, 0.95, 0.95, 1.0, 1.0, 1.0, 0.85]
    table = doc.add_table(rows=1 + len(pages), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_full_width(table)

    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "0F172A")
        set_cell_borders(cell, "0F172A")
        v_align(cell)
        set_cell_width(cell, widths[i])
        set_cell_text(cell, label, size=8, bold=True, color=WHITE, align="center")

    last_module = None
    for ri, page in enumerate(pages):
        row = table.rows[ri + 1]
        module_label = page["module"] if page["module"] != last_module else ""
        last_module = page["module"]
        values = [module_label, page["page"]]
        status_values = [
            page["statuses"].get(key, "") for key, _label in TASK_COLUMNS
        ]
        for ci, value in enumerate(values + status_values):
            cell = row.cells[ci]
            fill = "F8FAFC" if ri % 2 == 0 else "FFFFFF"
            color = NAVY
            bold = False
            align = "left"
            size = 8.5
            if ci >= 2:
                fill = STATUS_BG.get(value, fill)
                color = STATUS_COLOR.get(value, NAVY)
                bold = True
                align = "center"
                size = 8
                value = STATUS_SHORT.get(value, value)
            shade_cell(cell, fill)
            set_cell_borders(cell, "E2E8F0")
            v_align(cell)
            set_cell_width(cell, widths[ci])
            set_cell_text(
                cell,
                value,
                size=size,
                bold=bold or ci == 0,
                color=color,
                align=align,
            )

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_timeline_table(doc):
    rows = [
        (
            "1–6 September",
            "Complete remaining implementation (API client and API wiring).",
        ),
        (
            "7–12 September",
            "Staging environment preparation, and QA and UAT.",
        ),
        (
            "14–18 September",
            "Production readiness, go-live, and handover.",
        ),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_full_width(table)
    widths = [2.4, 7.6]
    headers = ["Week", "Focus"]
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "0F172A")
        set_cell_borders(cell, "0F172A")
        v_align(cell)
        set_cell_width(cell, widths[i])
        set_cell_text(cell, label, size=9, bold=True, color=WHITE)
    for ri, (week, focus) in enumerate(rows):
        fill = "F8FAFC" if ri % 2 == 0 else "FFFFFF"
        for ci, value in enumerate((week, focus)):
            cell = table.rows[ri + 1].cells[ci]
            shade_cell(cell, fill)
            set_cell_borders(cell, "E2E8F0")
            v_align(cell)
            set_cell_width(cell, widths[ci])
            set_cell_text(cell, value, size=10, bold=ci == 0)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def build():
    checklist = load_checklist()
    rows, counts = page_rows(checklist.ROLES)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = (
        section.page_height,
        section.page_width,
    )
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    title = doc.add_paragraph()
    run = title.add_run("eSariSari Float Management — Implementation status")
    set_run(run, size=20, bold=True, color=NAVY)
    title.paragraph_format.space_after = Pt(2)

    add_para(
        doc,
        "1 September 2026  ·  Client summary",
        size=11,
        color=MUTED,
        space_after=8,
    )
    add_para(
        doc,
        "Live demo screens are in place for the modules below. "
        "Remaining work follows the delivery timeline.",
        size=11,
        space_after=8,
    )

    heading(doc, "Delivery timeline", 1)
    add_timeline_table(doc)

    add_para(
        doc,
        f"{len(rows)} pages  ·  {counts['Completed']} completed  ·  "
        f"{counts['In progress']} in progress  ·  {counts['To do']} to do.",
        size=11,
        bold=True,
        space_after=4,
    )
    add_para(
        doc,
        "Completed = live in the current demo.  "
        "In progress = remaining implementation (API client and API wiring).  "
        "To do = QA/UAT, or not live for that role.  "
        "Pages marked (not live) are greyed in the app and not part of this release.",
        size=10,
        color=MUTED,
        space_after=10,
    )

    grouped = {}
    order = []
    for row in rows:
        if row["role"] not in grouped:
            grouped[row["role"]] = []
            order.append(row["role"])
        grouped[row["role"]].append(row)

    for role_key in order:
        heading(doc, ROLE_LABEL.get(role_key, role_key), 1)
        add_para(
            doc,
            ROLE_BLURB.get(role_key, ""),
            size=10,
            italic=True,
            color=MUTED,
            space_after=6,
        )
        add_role_table(doc, grouped[role_key])

    add_para(
        doc,
        "Dashboard / Franchisees / Retailers marked (not live) stay out of this "
        "release unless requested.",
        size=10,
        color=MUTED,
        space_after=0,
    )

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    print(f"Pages: {len(rows)}  ·  {counts}")


if __name__ == "__main__":
    build()
